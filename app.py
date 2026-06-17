from __future__ import annotations

import io
import hashlib
import hmac
import json
import os
import re
import time
import socket
import urllib.error
import urllib.parse
import urllib.request
from datetime import datetime
from typing import Optional

import streamlit as st
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# =============================================================================
# Page config
# =============================================================================
st.set_page_config(
    page_title="AI Novel Studio",
    page_icon="📖",
    layout="wide",
    initial_sidebar_state="expanded",
)

st.markdown(
    """
<style>
    .stage-header { font-size: 1.35rem; font-weight: 700; margin-bottom: 0.25rem; }
    .stage-desc { color: #777; font-size: 0.92rem; margin-bottom: 1.2rem; }
    .small-note { color: #777; font-size: 0.86rem; }
    .project-card { border: 1px solid rgba(128,128,128,.25); border-radius: 12px; padding: 14px; margin: 8px 0; }
    .accent-line { height: 3px; background: linear-gradient(90deg, #1D9E75, transparent); border-radius: 3px; margin: 16px 0; }
    .version-box { border-left: 4px solid #1D9E75; padding-left: 12px; margin: 8px 0; }
    .stProgress > div > div { background-color: #1D9E75; }
    div[data-testid="stSidebarContent"] { padding-top: 1.2rem; }
    .mobile-card { border: 1px solid rgba(128,128,128,.25); border-radius: 14px; padding: 12px; margin: 8px 0 14px 0; background: rgba(128,128,128,.04); }
    .mobile-kpi { font-size: 0.86rem; color: #666; margin-bottom: 0.25rem; }
    .mobile-title { font-size: 1.1rem; font-weight: 700; margin-bottom: 0.35rem; }
    @media (max-width: 760px) {
        .block-container { padding-left: 0.85rem; padding-right: 0.85rem; padding-top: 0.85rem; }
        .stage-header { font-size: 1.12rem; }
        .stage-desc { font-size: 0.86rem; margin-bottom: 0.75rem; }
        div[data-testid="stHorizontalBlock"] { gap: 0.35rem; }
        textarea { font-size: 16px !important; }
        input { font-size: 16px !important; }
    }
</style>
""",
    unsafe_allow_html=True,
)

# =============================================================================
# AI provider config: Gemini first, BigModel/Zhipu second
# =============================================================================
DEFAULT_GEMINI_MODEL = os.getenv("GEMINI_MODEL", "gemini-2.5-flash")
GEMINI_FALLBACK_MODELS = [
    DEFAULT_GEMINI_MODEL,
    "gemini-2.5-flash-lite",
    "gemini-2.0-flash",
    "gemini-1.5-flash",
]
DEFAULT_GLM_MODEL = os.getenv("BIGMODEL_MODEL", "glm-4-flash-250414")


def _read_secret_or_env(secret_name: str, env_name: str | None = None) -> Optional[str]:
    env_name = env_name or secret_name
    try:
        value = st.secrets.get(secret_name, None)
        if value:
            return str(value).strip()
    except Exception:
        pass
    value = os.getenv(env_name)
    return value.strip() if value else None


def get_gemini_api_key() -> Optional[str]:
    return _read_secret_or_env("GEMINI_API_KEY")


def get_bigmodel_api_key() -> Optional[str]:
    return (
        _read_secret_or_env("BIGMODEL_API_KEY")
        or _read_secret_or_env("ZHIPU_API_KEY")
        or _read_secret_or_env("ZAI_API_KEY")
    )


def _extract_gemini_text(result: dict) -> str:
    candidates = result.get("candidates", [])
    if not candidates:
        raise RuntimeError(f"No candidates returned from Gemini: {json.dumps(result, ensure_ascii=False)[:800]}")
    parts = candidates[0].get("content", {}).get("parts", [])
    texts = [p.get("text", "") for p in parts if p.get("text")]
    if texts:
        return "\n".join(texts).strip()
    raise RuntimeError(f"No text returned from Gemini: {json.dumps(result, ensure_ascii=False)[:800]}")


class AIProviderError(RuntimeError):
    """Provider error with enough metadata for smarter retry/fallback decisions."""

    def __init__(self, message: str, *, status_code: int | None = None, retryable: bool = False, fatal: bool = False):
        super().__init__(message)
        self.status_code = status_code
        self.retryable = retryable
        self.fatal = fatal


def _provider_status(message: str) -> None:
    """Best-effort status breadcrumb for the sidebar Last action box."""
    try:
        st.session_state["_last_action"] = message
    except Exception:
        pass


def _gemini_error_from_http(exc: urllib.error.HTTPError, clean_model: str) -> AIProviderError:
    error_body = exc.read().decode("utf-8", errors="ignore")
    code = int(getattr(exc, "code", 0) or 0)
    retryable = code == 429 or 500 <= code <= 599
    fatal = code in {401, 403}
    message = f"HTTP {code} for {clean_model}: {error_body[:800]}"
    return AIProviderError(message, status_code=code, retryable=retryable, fatal=fatal)


def _call_one_gemini_model(api_key: str, prompt: str, model_name: str, max_tokens: int) -> str:
    clean_model = model_name.replace("models/", "").strip()
    model = urllib.parse.quote(clean_model, safe="")
    url = f"https://generativelanguage.googleapis.com/v1beta/models/{model}:generateContent?key={urllib.parse.quote(api_key)}"
    payload = {
        "contents": [{"role": "user", "parts": [{"text": prompt}]}],
        "generationConfig": {
            "temperature": 0.82,
            "topP": 0.95,
            "maxOutputTokens": max_tokens,
        },
    }
    request = urllib.request.Request(
        url,
        data=json.dumps(payload).encode("utf-8"),
        headers={"Content-Type": "application/json"},
        method="POST",
    )
    timeout_seconds = float(os.getenv("AI_PROVIDER_TIMEOUT_SECONDS", "60"))
    try:
        with urllib.request.urlopen(request, timeout=timeout_seconds) as response:
            body = response.read().decode("utf-8")
    except urllib.error.HTTPError as exc:
        raise _gemini_error_from_http(exc, clean_model) from exc
    except (urllib.error.URLError, TimeoutError, socket.timeout) as exc:
        raise AIProviderError(
            f"Temporary network error calling Gemini {clean_model}: {exc}",
            retryable=True,
        ) from exc
    try:
        return _extract_gemini_text(json.loads(body))
    except Exception as exc:
        raise AIProviderError(f"Invalid Gemini response from {clean_model}: {exc}", retryable=False) from exc


def list_generate_content_models(api_key: str) -> list[str]:
    url = f"https://generativelanguage.googleapis.com/v1beta/models?key={urllib.parse.quote(api_key)}"
    request = urllib.request.Request(url, method="GET")
    try:
        with urllib.request.urlopen(request, timeout=30) as response:
            body = response.read().decode("utf-8")
        payload = json.loads(body)
    except Exception:
        return []
    models = []
    for model in payload.get("models", []):
        methods = model.get("supportedGenerationMethods", [])
        name = model.get("name", "").replace("models/", "")
        if name and "generateContent" in methods:
            models.append(name)
    return models


def call_gemini_rest(api_key: str, prompt: str, max_tokens: int = 3000) -> str:
    """Call Gemini with retry-aware fallback.

    Retryable errors (429, 5xx, timeout/network) get 1-2 short backoff retries on
    the same model. Non-retryable errors (bad request, model not found) switch to
    the next model immediately. Auth errors stop Gemini attempts quickly so the
    user does not wait through every fallback model.
    """
    candidate_models: list[str] = []
    for m in GEMINI_FALLBACK_MODELS:
        clean = m.replace("models/", "").strip()
        if clean and clean not in candidate_models:
            candidate_models.append(clean)

    max_dynamic = int(os.getenv("GEMINI_DYNAMIC_MODEL_LIMIT", "3"))
    available = list_generate_content_models(api_key)[:max_dynamic]
    for model in available:
        clean = model.replace("models/", "").strip()
        if clean and clean not in candidate_models:
            candidate_models.append(clean)

    max_retries = int(os.getenv("AI_RETRYABLE_RETRIES", "2"))
    base_delay = float(os.getenv("AI_RETRY_BASE_SECONDS", "1.5"))
    errors: list[str] = []

    for model in candidate_models:
        attempt = 0
        while True:
            try:
                _provider_status(f"Calling Gemini {model} (attempt {attempt + 1})...")
                return _call_one_gemini_model(api_key, prompt, model, max_tokens)
            except AIProviderError as exc:
                errors.append(str(exc))
                if exc.fatal:
                    raise RuntimeError("Gemini authentication/permission failed. Check API key or quota. " + str(exc)) from exc
                if exc.retryable and attempt < max_retries:
                    delay = min(base_delay * (2 ** attempt), 8.0)
                    _provider_status(f"Gemini {model} temporary error; retry {attempt + 1}/{max_retries} after {delay:.1f}s")
                    time.sleep(delay)
                    attempt += 1
                    continue
                # Non-retryable or retry budget exhausted: move to next model.
                _provider_status(f"Switching Gemini model after error on {model}: {str(exc)[:180]}")
                break
            except Exception as exc:
                errors.append(str(exc))
                _provider_status(f"Switching Gemini model after unexpected error on {model}: {str(exc)[:180]}")
                break

    raise RuntimeError("Could not call Gemini. Last errors: " + " | ".join(errors[-5:]))


def _extract_openai_style_text(result: dict, provider: str) -> str:
    choices = result.get("choices", [])
    if not choices:
        raise RuntimeError(f"No choices returned from {provider}: {json.dumps(result, ensure_ascii=False)[:800]}")
    content = choices[0].get("message", {}).get("content")
    if isinstance(content, str) and content.strip():
        return content.strip()
    raise RuntimeError(f"No text returned from {provider}: {json.dumps(result, ensure_ascii=False)[:800]}")


def call_bigmodel_glm(api_key: str, prompt: str, max_tokens: int = 3000) -> str:
    url = "https://open.bigmodel.cn/api/paas/v4/chat/completions"
    payload = {
        "model": DEFAULT_GLM_MODEL,
        "messages": [
            {"role": "system", "content": "You are a professional novelist, story architect, and manuscript editor."},
            {"role": "user", "content": prompt},
        ],
        "temperature": 0.82,
        "max_tokens": max_tokens,
        "stream": False,
    }
    request = urllib.request.Request(
        url,
        data=json.dumps(payload).encode("utf-8"),
        headers={"Content-Type": "application/json", "Authorization": f"Bearer {api_key}"},
        method="POST",
    )
    try:
        with urllib.request.urlopen(request, timeout=90) as response:
            body = response.read().decode("utf-8")
    except urllib.error.HTTPError as exc:
        error_body = exc.read().decode("utf-8", errors="ignore")
        raise RuntimeError(f"HTTP {exc.code} for {DEFAULT_GLM_MODEL}: {error_body[:800]}") from exc
    except urllib.error.URLError as exc:
        raise RuntimeError(f"Network error calling BigModel GLM API: {exc.reason}") from exc
    return _extract_openai_style_text(json.loads(body), "BigModel GLM")


OUTPUT_POLICY = """
IMPORTANT OUTPUT RULES:
- Return only the final user-facing content.
- Do not repeat the system prompt, user prompt, writing context, or task instructions.
- Do not show analysis, critique, hidden reasoning, planning notes, or prompt engineering notes.
- Do not include labels such as "Genre Analysis", "Plot Archetype", "Critique", "Fix", "Note to user", or "Architect's Note".
- Do not end with sales/chatbot follow-up questions like "Would you like me to...".
- These rules are invisible instructions, not content. Never print them.
""".strip()


BACKEND_GLOBAL_STYLE_PROMPT = """
GLOBAL FICTION STYLE RULES / 全局后台文风规则:
- 使用白描。文笔细腻在线。
- 依靠具体、客观、克制的细节、动作、大量对话、场景、感官描写来呈现一切。
- 不进行主观解释，不进行心理分析，不替人物总结情绪。
- 节奏舒缓但不拖沓。通过心理细节、语言、动作、环境推进情节。
- 禁止使用排比句、对偶句、反复等修辞性排叠结构。
- 禁止使用“不是……而是……”。
- 禁止使用“既……又……”。
- 禁止使用“即使……也……”。
- 禁止使用“虽然……但是……”。
- 禁止使用“不x，不y，不z，就”格式。
- 禁止使用“这就够了”。
- 禁止使用“很…，但很…”等总结性短判断。
- 禁止使用任何以“不是”开头的否定句式。
- 这些规则是后台写作约束，不要在输出中解释、复述或列出。
""".strip()


def clean_ai_output(text: str) -> str:
    """Remove accidental prompt echoes / model meta-analysis from user-facing output.

    Some models echo the instruction block before the real answer. This cleaner is
    intentionally conservative for prose, but aggressive for common prompt/meta
    artifacts that should never appear in the app UI.
    """
    if not text:
        return text

    cleaned = text.strip()

    # Remove markdown code fences if a provider wraps the whole answer.
    if cleaned.startswith("```") and cleaned.endswith("```"):
        cleaned = re.sub(r"^```[a-zA-Z0-9_-]*\s*", "", cleaned)
        cleaned = re.sub(r"\s*```$", "", cleaned).strip()

    # If the model produced a full prompt/analysis dump followed by the real
    # deliverable, prefer the LAST real Chinese Story Bible title.
    story_title_matches = list(re.finditer(r"《[^》]{1,80}》[^\n]{0,40}(?:故事设定集|设定集|Story Bible)", cleaned, flags=re.IGNORECASE))
    if story_title_matches:
        cleaned = cleaned[story_title_matches[-1].start():].strip()

    # Otherwise, if there are multiple deliverable-looking openings after a long
    # meta preface, cut to the last strong opening instead of the first echoed one.
    meta_markers = [
        "Genre Analysis", "Plot Archetype", "Female-oriented", "Web-novel",
        "Core Premise: Needs", "Critique:", "Fix:", "Since you have provided",
        "You are a", "Senior Story Architect", "Writing context:", "USER REQUEST:",
        "Check against all constraints", "Return *only*", "Return only the finished",
    ]
    if any(marker in cleaned[:4000] for marker in meta_markers):
        start_patterns = [
            r"^#\s*[^\n]+",
            r"^##\s*Core premise",
            r"^##\s*核心",
            r"^1\.\s*Core Premise",
            r"^1[、.]\s*核心",
            r"^Core premise\s*\(",
            r"^Core Premise\s*\(",
            r"^核心[前提设定]",
            r"^Logline\s*:",
            r"^第\s*[0-9一二两三四五六七八九十百]+\s*[章节章回]",
            r"^Chapter\s+\d+",
        ]
        starts = []
        for pat in start_patterns:
            for m in re.finditer(pat, cleaned, flags=re.IGNORECASE | re.MULTILINE):
                starts.append(m.start())
        # Keep a late opening if one exists after the meta dump.
        late_starts = [x for x in starts if x > 600]
        if late_starts:
            cleaned = cleaned[min(late_starts):].strip()
        elif starts and starts[0] > 0:
            cleaned = cleaned[starts[0]:].strip()

    # Remove output-policy / prompt-instruction lines that the model may echo.
    banned_line_patterns = [
        r"^\s*[*•-]\s*Return \*?only\*?.*$",
        r"^\s*[*•-]\s*No analysis.*$",
        r"^\s*[*•-]\s*No follow-up questions.*$",
        r"^\s*[*•-]\s*Start directly.*$",
        r"^\s*[*•-]\s*Use natural Chinese prose.*$",
        r"^\s*[*•-]\s*Follow \"?Shuangwen.*$",
        r"^\s*[*•-]\s*Adhere to banned phrases.*$",
        r"^\s*Check against all constraints.*$",
        r"^\s*No \"Genre Analysis\".*$",
        r"^\s*Only the content\.?\s*$",
        r"^\s*Language is Chinese\.?\s*$",
        r"^\s*Banned phrases included.*$",
        r"^\s*Title:\s*.*$",
    ]
    kept_lines = []
    for line in cleaned.splitlines():
        if any(re.match(pat, line, flags=re.IGNORECASE) for pat in banned_line_patterns):
            continue
        kept_lines.append(line)
    cleaned = "\n".join(kept_lines).strip()

    # Remove common trailing meta notes / follow-up offers.
    trailing_patterns = [
        r"\n*Architect's Note to User:.*$",
        r"\n*Note to User:.*$",
        r"\n*作者提示[:：].*$",
        r"\n*Would you like me to .*$",
        r"\n*如果你愿意，我可以.*$",
        r"\n*你可以继续让我.*$",
    ]
    for pat in trailing_patterns:
        cleaned = re.sub(pat, "", cleaned, flags=re.IGNORECASE | re.DOTALL).strip()

    return cleaned





def clean_chapter_draft_output(text: str) -> str:
    """Keep only final chapter prose when a model leaks planning notes.

    Draft Writer outputs are especially vulnerable to echoed beat sheets,
    self-correction notes, and lines such as "Let's write". This cleaner
    removes those artifacts before anything is displayed or saved.
    """
    cleaned = clean_ai_output(text or "").strip()
    if not cleaned:
        return cleaned

    # Strong cut: many leaked drafts place the real prose after "Let's write"
    # or a similar transition marker. Keep only what follows the last marker.
    transition_patterns = [
        r"(?i)let['’]?s write\.?",
        r"(?i)now write the chapter\.?",
        r"正文如下\s*[:：]",
        r"开始正文\s*[:：]?",
        r"下面是正文\s*[:：]?",
    ]
    last_end = -1
    for pat in transition_patterns:
        for m in re.finditer(pat, cleaned):
            last_end = max(last_end, m.end())
    if last_end >= 0:
        cleaned = cleaned[last_end:].lstrip(" \n\t:：-—").strip()

    meta_line_patterns = [
        r"^\s*[*•\-]?\s*\*?Opening\*?\s*[:：]",
        r"^\s*[*•\-]?\s*\*?Transition\*?\s*[:：]",
        r"^\s*[*•\-]?\s*\*?Conflict\*?\s*[:：]",
        r"^\s*[*•\-]?\s*\*?The Move\*?\s*[:：]",
        r"^\s*[*•\-]?\s*\*?The Aftermath\*?\s*[:：]",
        r"^\s*[*•\-]?\s*\*?The Selection\*?\s*[:：]",
        r"^\s*[*•\-]?\s*\*?Selection\*?\s*[:：]",
        r"^\s*[*•\-]?\s*\*?Death Scene\*?\s*[:：]",
        r"^\s*[*•\-]?\s*\*?Death/Betrayal\*?",
        r"^\s*[*•\-]?\s*\*?Rebirth\*?\s*[:：]",
        r"^\s*[*•\-]?\s*\*?Dialogue\*?\s*[:：]",
        r"^\s*[*•\-]?\s*\*?Internal Monologue\*?\s*[:：]",
        r"^\s*[*•\-]?\s*\*?Action\*?\s*[:：]",
        r"^\s*[*•\-]?\s*\*?Ending\*?\s*[:：]",
        r"^\s*[*•\-]?\s*\*?Check\*?\s*[:：]",
        r"^\s*[*•\-]?\s*\*?Initial thought\*?\s*[:：]",
        r"^\s*[*•\-]?\s*\*?Correction\*?\s*[:：]",
        r"^\s*[*•\-]?\s*\*?Language\*?\s*[:：]",
        r"^\s*[*•\-]?\s*\*?Style\*?\s*[:：]",
        r"^\s*[*•\-]?\s*\*?Structure\*?\s*[:：]",
        r"^\s*\*?\(?Self-Correction",
        r"^\s*\*?Drafting",
        r"^\s*\*?Final Polish",
        r"^\s*\*?Wait, the prompt",
        r"^\s*Return only the final",
        r"^\s*No \"deep breath\"",
        r"^\s*Xianxia,",
        r"^\s*Female-oriented",
        r"^\s*Modern Webnovel",
    ]

    def is_meta_line(line: str) -> bool:
        return any(re.match(pat, line, flags=re.IGNORECASE) for pat in meta_line_patterns)

    lines = cleaned.splitlines()
    last_meta_idx = -1
    # Look only near the beginning so we do not remove intentional later prose.
    for i, line in enumerate(lines[:120]):
        if is_meta_line(line):
            last_meta_idx = i

    if last_meta_idx >= 0:
        for j in range(last_meta_idx + 1, len(lines)):
            candidate = re.sub(r"^\s*[*•\-]+\s*", "", lines[j]).strip()
            candidate = candidate.strip("* ")
            if not candidate or is_meta_line(candidate):
                continue
            # A real Chinese webnovel paragraph usually contains several CJK
            # characters and is not a planning label.
            cjk_count = len(re.findall(r"[\u4e00-\u9fff]", candidate))
            if cjk_count >= 8:
                cleaned = "\n".join(lines[j:]).strip()
                break

    # Drop any remaining meta/check lines that slipped into the top.
    kept = []
    for line in cleaned.splitlines():
        if is_meta_line(line):
            continue
        if re.match(r"^\s*[*•\-]\s*\*?(Opening|Transition|Conflict|Selection|Check|Language|Style|Structure)\*?\s*[:：]", line, flags=re.IGNORECASE):
            continue
        kept.append(line)
    cleaned = "\n".join(kept).strip()

    # If a chapter title exists later after a leaked preface, cut to it.
    chapter_starts = list(re.finditer(r"^(?:第\s*[0-9一二两三四五六七八九十百]+\s*[章节章回]|Chapter\s+\d+)", cleaned, flags=re.IGNORECASE | re.MULTILINE))
    if chapter_starts:
        first = chapter_starts[0].start()
        if first > 300:
            cleaned = cleaned[first:].strip()

    return cleaned



def clean_story_bible_output(text: str, title: str = "") -> str:
    """Extra-aggressive cleanup for Story Bible outputs.

    Some models expose their planning/prompt notes before the actual bible. For this
    stage, we only want the finished deliverable shown to the writer.
    """
    cleaned = clean_ai_output(text or "").strip()
    if not cleaned:
        return cleaned

    # If the model dumped planning then later started the real Chinese bible at
    # 核心前提, keep the LAST real section start. This removes earlier prompt-like
    # bullets such as "* **## 核心前提 ..." from the model's analysis.
    core_starts = []
    core_pat = r"^(?:\s*[*•\-]\s*)?(?:\*\*)?(?:#+\s*)?核心前提\b|^(?:\s*[*•\-]\s*)?(?:\*\*)?(?:#+\s*)?核心前提[（(]"
    for m in re.finditer(core_pat, cleaned, flags=re.MULTILINE):
        core_starts.append(m.start())
    meta_markers = [
        "Senior Novel Setting Architect", "Story Setting Bible", "Only the final content",
        "No Architect", "Specific Chapter Headings", "Theme:", "Setting:",
        "Protagonist", "Conflict:", "Power System:", "Self-Correction",
        "Check Chapter Headings", "Check Content", "Drafting the content",
        "Mental translation", "Key concept", "Return only", "No analysis",
    ]
    if core_starts and (len(core_starts) > 1 or any(x in cleaned[:7000] for x in meta_markers)):
        cleaned = cleaned[core_starts[-1]:].strip()

    # Remove common markdown wrappers from the first line after cutting.
    cleaned = re.sub(r"^\s*[*•\-]\s*", "", cleaned).strip()
    cleaned = re.sub(r"^\*\*(.*?)\*\*", r"\1", cleaned, count=1, flags=re.DOTALL).strip()

    # Normalize the title: user-facing Story Bible should begin cleanly with
    # 《书名》故事设定集, not prompt notes or English labels.
    safe_title = (title or "My Novel").strip() or "My Novel"
    title_line = f"《{safe_title}》故事设定集"
    if not re.match(r"^《[^》]{1,80}》[^\n]{0,30}(故事设定集|设定集)", cleaned):
        cleaned = title_line + "\n\n" + cleaned

    # Delete any lingering prompt/instruction lines anywhere in the first screen.
    bad_line_pats = [
        r"^Senior .*Architect.*$",
        r"^A complete .*Story.*Bible.*$",
        r"^\s*[*•\-]\s*Only the final content.*$",
        r"^\s*[*•\-]\s*No .*$",
        r"^\s*[*•\-]\s*Specific Chapter Headings.*$",
        r"^\s*[*•\-]\s*Theme:.*$",
        r"^\s*[*•\-]\s*Setting:.*$",
        r"^\s*[*•\-]\s*Protagonist.*$",
        r"^\s*[*•\-]\s*Conflict:.*$",
        r"^\s*[*•\-]\s*Power System:.*$",
        r"^\s*\(?Self-Correction.*$",
        r"^\s*\*?Check .*",
        r"^\s*\*?Drafting the content.*$",
        r"^\s*\*?Key concept.*$",
    ]
    kept=[]
    for line in cleaned.splitlines():
        if any(re.match(pat, line, flags=re.IGNORECASE) for pat in bad_line_pats):
            continue
        kept.append(line)
    cleaned="\n".join(kept).strip()

    # Ensure title remains first after line deletion.
    if not cleaned.startswith("《"):
        cleaned = title_line + "\n\n" + cleaned
    return cleaned


def call_ai(system: str, user: str, max_tokens: int = 3000) -> str:
    prompt = f"{system.strip()}\n\n{BACKEND_GLOBAL_STYLE_PROMPT}\n\n{OUTPUT_POLICY}\n\nUSER REQUEST:\n{user.strip()}"
    gemini_key = get_gemini_api_key()
    bigmodel_key = get_bigmodel_api_key()
    errors = []
    if gemini_key:
        try:
            return call_gemini_rest(gemini_key, prompt, max_tokens=max_tokens)
        except Exception as exc:
            errors.append(f"Gemini failed: {exc}")
    if bigmodel_key:
        try:
            return call_bigmodel_glm(bigmodel_key, prompt, max_tokens=max_tokens)
        except Exception as exc:
            errors.append(f"BigModel/Zhipu failed: {exc}")
    if not gemini_key and not bigmodel_key:
        raise RuntimeError("No AI key is configured. Add GEMINI_API_KEY or BIGMODEL_API_KEY in Streamlit secrets.")
    raise RuntimeError("AI providers failed. " + " | ".join(errors[-3:]))


def generate_text(system: str, user: str, placeholder, max_tokens: int = 3000, cleaner=None, task_name: str = "AI generation") -> str:
    """Call AI with basic generation locks and raw-output backup.

    Streamlit runs top-to-bottom on every interaction. This wrapper protects the
    app from losing long generations by saving raw and cleaned output to
    session_state before rendering it in the placeholder.
    """
    if st.session_state.get("ai_running", False):
        st.warning("AI is already generating. Please wait until the current task finishes.")
        return ""
    st.session_state.ai_running = True
    st.session_state.ai_task_name = task_name
    st.session_state.ai_started_at = datetime.now().isoformat(timespec="seconds")
    try:
        raw_text = call_ai(system, user, max_tokens=max_tokens)
        st.session_state["_last_ai_raw_output"] = raw_text
        st.session_state["_last_ai_task_name"] = task_name
        st.session_state["_last_ai_raw_at"] = datetime.now().isoformat(timespec="seconds")
        text = clean_ai_output(raw_text)
        if cleaner is not None:
            text = cleaner(text)
        # Save cleaned output before drawing, so a UI refresh does not lose it.
        st.session_state["_last_ai_clean_output"] = text
    except Exception as exc:
        st.error(f"AI call failed: {exc}")
        return ""
    finally:
        st.session_state.ai_running = False
        st.session_state.ai_task_name = ""
    shown = ""
    for chunk in re.split(r"(\n\n+)", text):
        shown += chunk
        placeholder.markdown(shown + "▌")
        time.sleep(0.005)
    placeholder.markdown(text)
    return text

# =============================================================================
# State and project helpers
# =============================================================================
PROJECT_FIELDS = [
    "stage", "project_name", "novel_title", "language", "chinese_style", "audience_channel",
    "publishing_platforms", "submission_types", "platform_goals", "custom_platform_note",
    "genre", "genre_tags", "custom_genre_tags", "target_length",
    "story_idea", "style_note", "global_instruction", "banned_phrases", "story_bible", "characters",
    "worldbuilding", "timeline", "glossary", "outline", "chapters", "chapter_beats", "chapter_drafts",
    "chapter_versions", "pending_revisions", "local_refinements", "full_draft", "logic_report", "format_report", "polished_draft",
    "quality_reports", "market_reviews", "revision_notes",
]

DEFAULTS = {
    "stage": 1,
    "project_name": "my_novel_project",
    "novel_title": "My Novel",
    "language": "Chinese 中文",
    "chinese_style": "现代网文",
    "audience_channel": "女频",
    "publishing_platforms": ["晋江文学城", "番茄小说"],
    "submission_types": ["长篇连载", "签约向商业文"],
    "platform_goals": ["前三章留存", "签约通过率", "追读率"],
    "custom_platform_note": "",
    "genre": "Xianxia / 仙侠",
    "genre_tags": ["重生", "复仇", "大女主"],
    "custom_genre_tags": "",
    "target_length": "Novel (~18 chapters)",
    "story_idea": "",
    "style_note": "",
    "global_instruction": "",
    "banned_phrases": "她深吸一口气\n空气仿佛凝固了\n命运的齿轮开始转动\n这一切才刚刚开始\n眼中闪过一丝复杂",
    "story_bible": "",
    "characters": "",
    "worldbuilding": "",
    "timeline": "",
    "glossary": "",
    "outline": "",
    "chapters": [],
    "chapter_beats": {},
    "chapter_drafts": {},
    "chapter_versions": {},
    "pending_revisions": {},
    "local_refinements": {},
    "full_draft": "",
    "logic_report": "",
    "format_report": "",
    "polished_draft": "",
    "quality_reports": {},
    "market_reviews": {},
    "revision_notes": {},
}

for key, value in DEFAULTS.items():
    if key not in st.session_state:
        st.session_state[key] = value

# =============================================================================
# Stability helpers: dirty state, AI/save locks, login token
# =============================================================================
STABILITY_DEFAULTS = {
    "ai_running": False,
    "ai_task_name": "",
    "ai_started_at": "",
    "cloud_save_running": False,
    "cloud_last_save_status": "idle",
    "cloud_last_save_at": "",
    "cloud_last_saved_hash": "",
    "cloud_last_summary": "",
    "project_dirty": False,
    "pending_ai_results": {},
    "batch_revision_proposals": {},
    "cloud_snapshot_status": "",
    "view_mode": "Auto",
}
for key, value in STABILITY_DEFAULTS.items():
    if key not in st.session_state:
        st.session_state[key] = value


def mark_dirty() -> None:
    st.session_state.project_dirty = True


def _json_hash(data: object) -> str:
    return hashlib.sha256(json.dumps(data, ensure_ascii=False, sort_keys=True, default=str).encode("utf-8")).hexdigest()


def current_project_hash() -> str:
    return _json_hash(project_data())


def mark_saved(stamp: str = "") -> None:
    st.session_state.cloud_last_saved_hash = current_project_hash()
    st.session_state.project_dirty = False
    st.session_state.cloud_last_save_at = stamp or datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    st.session_state.cloud_last_save_status = "success"


def selected_value_is_valid(key: str, options: list, default):
    if st.session_state.get(key) not in options:
        st.session_state[key] = default


def selected_list_is_valid(key: str, options: list, default: list | None = None):
    current = st.session_state.get(key, default or []) or []
    st.session_state[key] = [x for x in current if x in options]
    if not st.session_state[key] and default:
        st.session_state[key] = [x for x in default if x in options]


def auth_secret() -> str:
    return (
        _read_secret_or_env("AUTH_TOKEN_SECRET")
        or _read_secret_or_env("SUPABASE_KEY")
        or _read_secret_or_env("SUPABASE_ANON_KEY")
        or "ai-novel-studio-local-dev-secret"
    )


def make_auth_token(username: str, days: int = 30) -> str:
    exp = int(time.time()) + days * 86400
    body = f"{username}:{exp}"
    sig = hmac.new(auth_secret().encode("utf-8"), body.encode("utf-8"), hashlib.sha256).hexdigest()
    return f"{body}:{sig}"


def verify_auth_token(token: str) -> str | None:
    try:
        username, exp_s, sig = token.split(":", 2)
        if int(exp_s) < int(time.time()):
            return None
        body = f"{username}:{exp_s}"
        expected = hmac.new(auth_secret().encode("utf-8"), body.encode("utf-8"), hashlib.sha256).hexdigest()
        return username if hmac.compare_digest(sig, expected) else None
    except Exception:
        return None


def restore_login_from_query_token() -> None:
    if st.session_state.get("auth_ok"):
        return
    try:
        token = st.query_params.get("auth", "")
    except Exception:
        token = ""
    username = verify_auth_token(token) if token else None
    if username:
        st.session_state.auth_ok = True
        st.session_state.auth_user = username


def clear_login_query_token() -> None:
    try:
        if "auth" in st.query_params:
            del st.query_params["auth"]
    except Exception:
        pass


def render_status_bar() -> None:
    user = st.session_state.get("auth_user", "") or "not logged in"
    dirty = "Unsaved changes" if st.session_state.get("project_dirty") else "Saved"
    ai = f"AI: {st.session_state.get('ai_task_name')}" if st.session_state.get("ai_running") else "AI: idle"
    cloud = st.session_state.get("cloud_last_save_at") or "not saved this session"
    st.caption(f"🔒 {user} | ☁️ {dirty} | Last save: {cloud} | {ai}")


def story_bible_is_incomplete(text: str) -> bool:
    if not text or len(text.strip()) < 1200:
        return True
    required = ["## 核心前提", "## 主要角色", "## 时间线", "## 必须避免"]
    return not all(x in text for x in required)


def render_formatted_outline() -> None:
    chapters = st.session_state.get("chapters", []) or parse_chapters(st.session_state.get("outline", ""))
    if not chapters:
        st.info("还没有识别到章节。可以查看原始大纲，或把章节格式改成：第1章：标题 — 摘要。")
        return
    for ch in chapters:
        num = ch.get("num", "")
        title = ch.get("title", "未命名章节")
        summary = ch.get("summary", "")
        st.markdown(f"### 第{num}章：{title}")
        st.markdown(summary or "_无摘要_")
        st.divider()


def normalize_widget_defaults() -> None:
    selected_value_is_valid("language", ["Chinese 中文", "English", "Bilingual 中英双语"], "Chinese 中文")
    selected_value_is_valid("audience_channel", ["女频", "男频", "通用/无固定频道"], "女频")
    selected_value_is_valid("chinese_style", ["现代网文", "古风/仙侠", "女频仙侠", "玄幻升级流", "悬疑推理", "都市情感", "轻小说", "宫斗权谋", "中式恐怖", "文学感", "自定义"], "现代网文")
    selected_value_is_valid("genre", ["Xianxia / 仙侠", "Wuxia / 武侠", "Fantasy", "Science Fiction", "Romance", "Thriller", "Mystery", "Historical Fiction", "Urban Fantasy / 都市异能", "Horror", "Literary Fiction", "Other"], "Xianxia / 仙侠")
    selected_value_is_valid("target_length", ["Short story (~5 chapters)", "Novella (~10 chapters)", "Novel (~18 chapters)", "Long web novel (~30 chapters)"], "Novel (~18 chapters)")


def normalize_cloud_widget_keys(key_prefix: str) -> None:
    # Keep sidebar/setup cloud widgets synced with global settings.
    auto_key = f"{key_prefix}_auto_save_enabled"
    if auto_key in st.session_state:
        st.session_state.cloud_auto_save_enabled = bool(st.session_state.get(auto_key))


def default_cloud_project_name() -> str:
    """Use Step 1 project name automatically for cloud save."""
    name = st.session_state.get("project_name") or st.session_state.get("novel_title") or "Untitled"
    return str(name).strip() or "Untitled"


def normalize_project_name_from_cloud_widget(key_prefix: str) -> str:
    save_as_key = f"{key_prefix}_save_as_name"
    override = str(st.session_state.get(save_as_key, "") or "").strip()
    return override or default_cloud_project_name()


def cloud_project_updated_at(username: str, project_name: str) -> str:
    try:
        for item in cloud_list_projects(username):
            if item.get("project_name") == project_name:
                return str(item.get("updated_at") or "")
    except Exception:
        return ""
    return ""


def cloud_save_project_guarded(username: str, project_name: str, force: bool = False, save_reason: str = "manual save") -> dict:
    if st.session_state.get("ai_running"):
        raise RuntimeError("AI is currently generating. Please wait before saving, so a partial/old result is not saved.")
    if st.session_state.get("cloud_save_running"):
        raise RuntimeError("Cloud save is already running.")
    cloud_updated = cloud_project_updated_at(username, project_name)
    last_local = st.session_state.get("cloud_last_save_at", "")
    if cloud_updated and last_local and cloud_updated[:19].replace("T", " ") > last_local and not force:
        raise RuntimeError("Cloud project changed after your last local save. Load cloud, save as a new project, or enable Force overwrite.")
    st.session_state.cloud_save_running = True
    st.session_state.cloud_last_save_status = "saving"
    try:
        saved = cloud_save_project(username, project_name)
        try:
            cloud_create_project_snapshot(username, project_name, save_reason=save_reason)
            st.session_state.cloud_snapshot_status = "snapshot saved"
        except Exception as snapshot_exc:
            st.session_state.cloud_snapshot_status = f"snapshot skipped/failed: {snapshot_exc}"
        updated = saved.get("updated_at", "") if isinstance(saved, dict) else ""
        stamp = updated[:19].replace("T", " ") if updated else datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        mark_saved(stamp)
        return saved
    finally:
        st.session_state.cloud_save_running = False


def normalize_loaded_project(data: dict) -> dict:
    normalized = dict(DEFAULTS)
    normalized.update({k: v for k, v in data.items() if k in DEFAULTS})
    for dict_key in ["chapter_beats", "chapter_drafts", "chapter_versions", "pending_revisions", "local_refinements", "quality_reports", "market_reviews", "revision_notes"]:
        value = normalized.get(dict_key, {})
        if isinstance(value, dict):
            normalized[dict_key] = {int(k) if str(k).isdigit() else k: v for k, v in value.items()}
    return normalized


def project_data() -> dict:
    data = {field: st.session_state.get(field, DEFAULTS.get(field)) for field in PROJECT_FIELDS}
    data["saved_at"] = datetime.now().isoformat(timespec="seconds")
    data["app_version"] = "novel-studio-project-v24-smart-retry-tests"
    return data


def load_project(data: dict) -> None:
    """Load a project/checkpoint JSON and make it immediately usable.

    Older checkpoints may have saved the outline but not the parsed
    ``chapters`` list. In that case, rebuild chapters from the outline so
    Chapter Planner / Draft Writer can continue instead of showing 0 chapters.
    """
    normalized = normalize_loaded_project(data)
    for key, value in normalized.items():
        st.session_state[key] = value

    note_parts: list[str] = []

    # Repair old or partial checkpoints: outline exists but chapters are empty.
    if (not st.session_state.get("chapters")) and str(st.session_state.get("outline", "")).strip():
        repaired = parse_chapters(st.session_state.outline)
        st.session_state.chapters = repaired
        if repaired:
            note_parts.append(f"已从大纲重新识别 {len(repaired)} 个章节。")

    # Stronger repair: some Stage 3 checkpoints were saved after Story Bible,
    # before the actual outline was generated. They contain story_bible/story_idea
    # but outline == "" and chapters == []. In that case, infer a temporary
    # outline from the Story Bible timeline so the author can continue, then
    # recommend regenerating a clean outline in Stage 3.
    if (not st.session_state.get("chapters")) and not str(st.session_state.get("outline", "")).strip():
        inferred = infer_chapters_from_story_bible_timeline(str(st.session_state.get("story_bible", "")))
        if inferred:
            st.session_state.chapters = inferred
            st.session_state.outline = build_recovered_outline_from_chapters(st.session_state.get("novel_title", "My Novel"), inferred)
            note_parts.append(
                f"该 JSON 没有保存正式大纲；已从 Story Bible 的时间线临时恢复 {len(inferred)} 个章节/阶段。建议在 Stage 3 重新生成或微调大纲。"
            )
        elif int(st.session_state.get("stage", 1) or 1) >= 3:
            st.session_state.stage = 3
            note_parts.append("该 JSON 保存时还没有大纲，因此暂时停在 Stage 3。请点击 Generate outline 生成大纲后继续。")

    # If the saved stage is after Outline but chapters still cannot be parsed,
    # send the user back to Outline so they can regenerate/repair the outline.
    if int(st.session_state.get("stage", 1) or 1) >= 4 and not st.session_state.get("chapters"):
        st.session_state.stage = 3
        note_parts.append("未识别到章节，已自动回到大纲阶段。")

    st.session_state["_json_repair_note"] = " ".join(p for p in note_parts if p).strip()
    rebuild_full_draft()
    st.session_state.project_dirty = False


def loaded_project_summary(data: dict | None = None) -> str:
    """Human-readable summary after JSON restore, based on repaired session state."""
    stage = int(st.session_state.get("stage", 1) or 1)
    chapters = st.session_state.get("chapters", []) or []
    beats = st.session_state.get("chapter_beats", {}) or {}
    drafts = st.session_state.get("chapter_drafts", {}) or {}
    title = st.session_state.get("novel_title", "Untitled")
    note = st.session_state.get("_json_repair_note", "")
    msg = (
        f"已恢复《{title}》：Stage {stage}，"
        f"章节 {len(chapters)} 个，Beat Sheet {len(beats)} 个，正文草稿 {len(drafts)} 章。"
    )
    if note:
        msg += " " + note
    return msg


# =============================================================================
# Supabase Cloud Save / Load
# =============================================================================
DEFAULT_SUPABASE_URL = "https://cfxtvflqqaqlrnkcmgig.supabase.co"


def get_supabase_url() -> Optional[str]:
    return _read_secret_or_env("SUPABASE_URL") or DEFAULT_SUPABASE_URL


def get_supabase_key() -> Optional[str]:
    return _read_secret_or_env("SUPABASE_KEY") or _read_secret_or_env("SUPABASE_ANON_KEY")


def cloud_enabled() -> bool:
    return bool(get_supabase_url() and get_supabase_key())


def _supabase_headers(extra: dict | None = None) -> dict:
    key = get_supabase_key()
    if not key:
        raise RuntimeError("Missing SUPABASE_KEY / SUPABASE_ANON_KEY in Streamlit secrets.")
    headers = {
        "apikey": key,
        "Authorization": f"Bearer {key}",
        "Content-Type": "application/json",
    }
    if extra:
        headers.update(extra)
    return headers


def _supabase_request(method: str, path: str, payload: dict | list | None = None, prefer: str | None = None) -> object:
    base = (get_supabase_url() or "").rstrip("/")
    if not base:
        raise RuntimeError("Missing SUPABASE_URL in Streamlit secrets.")
    url = f"{base}/rest/v1/{path.lstrip('/')}"
    headers = _supabase_headers({"Prefer": prefer} if prefer else None)
    data = None if payload is None else json.dumps(payload, ensure_ascii=False).encode("utf-8")
    req = urllib.request.Request(url, data=data, headers=headers, method=method)
    try:
        with urllib.request.urlopen(req, timeout=45) as response:
            body = response.read().decode("utf-8")
            if not body.strip():
                return None
            return json.loads(body)
    except urllib.error.HTTPError as exc:
        body = exc.read().decode("utf-8", errors="ignore")
        raise RuntimeError(f"Supabase HTTP {exc.code}: {body[:1000]}") from exc
    except urllib.error.URLError as exc:
        raise RuntimeError(f"Supabase network error: {exc.reason}") from exc


def cloud_list_projects(username: str) -> list[dict]:
    user_filter = urllib.parse.quote(username, safe="")
    path = (
        "novel_projects"
        "?select=id,username,project_name,novel_title,current_stage,updated_at"
        f"&username=eq.{user_filter}"
        "&order=updated_at.desc"
    )
    result = _supabase_request("GET", path)
    return result if isinstance(result, list) else []


def cloud_save_project(username: str, project_name: str | None = None) -> dict:
    name = (project_name or st.session_state.get("project_name") or st.session_state.get("novel_title") or "Untitled").strip()
    if not name:
        name = "Untitled"
    # Keep the Streamlit project name aligned with the cloud record name.
    st.session_state.project_name = name
    data = project_data()
    payload = {
        "username": username,
        "project_name": name,
        "novel_title": st.session_state.get("novel_title", name),
        "current_stage": int(st.session_state.get("stage", 1) or 1),
        "project_data": data,
    }
    path = "novel_projects?on_conflict=username,project_name"
    result = _supabase_request(
        "POST",
        path,
        payload,
        prefer="resolution=merge-duplicates,return=representation",
    )
    if isinstance(result, list) and result:
        return result[0]
    if isinstance(result, dict):
        return result
    return payload


def cloud_load_project(username: str, project_name: str) -> dict:
    user_filter = urllib.parse.quote(username, safe="")
    project_filter = urllib.parse.quote(project_name, safe="")
    path = (
        "novel_projects"
        "?select=project_data,project_name,novel_title,current_stage,updated_at"
        f"&username=eq.{user_filter}"
        f"&project_name=eq.{project_filter}"
        "&limit=1"
    )
    result = _supabase_request("GET", path)
    if not isinstance(result, list) or not result:
        raise RuntimeError(f"Cloud project not found: {project_name}")
    data = result[0].get("project_data")
    if not isinstance(data, dict):
        raise RuntimeError("Cloud project exists, but project_data is missing or invalid.")
    return data


def cloud_delete_project(username: str, project_name: str) -> None:
    user_filter = urllib.parse.quote(username, safe="")
    project_filter = urllib.parse.quote(project_name, safe="")
    path = f"novel_projects?username=eq.{user_filter}&project_name=eq.{project_filter}"
    _supabase_request("DELETE", path)


def cloud_create_project_snapshot(username: str, project_name: str, save_reason: str = "manual save") -> dict:
    data = project_data()
    payload = {
        "username": username,
        "project_name": project_name,
        "novel_title": st.session_state.get("novel_title", project_name),
        "current_stage": int(st.session_state.get("stage", 1) or 1),
        "save_reason": save_reason,
        "project_hash": _json_hash(data),
        "project_data": data,
    }
    result = _supabase_request("POST", "novel_project_snapshots", payload, prefer="return=representation")
    if isinstance(result, list) and result:
        return result[0]
    if isinstance(result, dict):
        return result
    return payload


def cloud_list_project_snapshots(username: str, project_name: str, limit: int = 30) -> list[dict]:
    user_filter = urllib.parse.quote(username, safe="")
    project_filter = urllib.parse.quote(project_name, safe="")
    path = (
        "novel_project_snapshots"
        "?select=id,username,project_name,novel_title,current_stage,save_reason,project_hash,saved_at"
        f"&username=eq.{user_filter}"
        f"&project_name=eq.{project_filter}"
        "&order=saved_at.desc"
        f"&limit={int(limit)}"
    )
    result = _supabase_request("GET", path)
    return result if isinstance(result, list) else []


def cloud_load_project_snapshot(snapshot_id: str) -> dict:
    snap_filter = urllib.parse.quote(str(snapshot_id), safe="")
    path = f"novel_project_snapshots?select=project_data&id=eq.{snap_filter}&limit=1"
    result = _supabase_request("GET", path)
    if not isinstance(result, list) or not result:
        raise RuntimeError("Snapshot not found.")
    data = result[0].get("project_data")
    if not isinstance(data, dict):
        raise RuntimeError("Snapshot exists, but project_data is missing or invalid.")
    return data


def cloud_auto_save(reason: str = "project update") -> bool:
    """Auto-save current project after meaningful generation events.

    This intentionally does not run on every Streamlit rerun or every text edit.
    It only runs when a generation/save action explicitly calls it.
    """
    if not st.session_state.get("cloud_auto_save_enabled", False):
        return False
    if st.session_state.get("ai_running"):
        st.session_state.cloud_last_summary = "Auto-save skipped while AI is generating."
        return False
    if not cloud_enabled():
        st.session_state.cloud_last_summary = "Auto-save skipped: Supabase is not configured."
        return False
    username = st.session_state.get("auth_user", "user") or "user"
    project_name = default_cloud_project_name()
    try:
        saved = cloud_save_project_guarded(username, project_name, force=True, save_reason=f"auto-save: {reason}")
        updated = saved.get("updated_at", "") if isinstance(saved, dict) else ""
        stamp = updated[:19].replace("T", " ") if updated else st.session_state.get("cloud_last_save_at", datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
        st.session_state.cloud_last_summary = f"Auto-saved '{project_name}' after {reason}. Last save: {stamp}"
        st.session_state["_last_action"] = st.session_state.cloud_last_summary
        try:
            st.toast("Auto-saved to cloud")
        except Exception:
            pass
        return True
    except Exception as exc:
        st.session_state.cloud_last_summary = f"Auto-save failed after {reason}: {exc}"
        st.session_state.cloud_last_save_status = "failed"
        st.session_state["_last_action"] = st.session_state.cloud_last_summary
        return False


def render_cloud_projects_panel(key_prefix: str = "cloud") -> None:
    username = st.session_state.get("auth_user", "user") or "user"
    with st.expander("☁️ Cloud Projects / 云端项目", expanded=False):
        st.caption("Save/load projects from Supabase. JSON download/upload remains available as backup.")
        if not cloud_enabled():
            st.warning("Cloud save is not configured yet. Add SUPABASE_URL and SUPABASE_KEY to Streamlit secrets.")
            st.code(
                'SUPABASE_URL = "https://cfxtvflqqaqlrnkcmgig.supabase.co"\nSUPABASE_KEY = "your-anon-public-key"',
                language="toml",
            )
            return

        auto_cloud_name = default_cloud_project_name()
        st.caption(f"Cloud project name is auto-filled from Step 1: **{auto_cloud_name}**")
        with st.expander("Save as new copy / 另存为新项目", expanded=False):
            st.text_input(
                "Optional new cloud project name",
                key=f"{key_prefix}_save_as_name",
                placeholder="Leave blank to use the Step 1 project file name",
            )
            st.caption("Leave blank for normal Save to cloud. Fill only when you intentionally want a new cloud copy.")
        if f"{key_prefix}_auto_save_enabled" not in st.session_state:
            st.session_state[f"{key_prefix}_auto_save_enabled"] = st.session_state.get("cloud_auto_save_enabled", False)
        st.checkbox(
            "Auto-save after generation / 生成后自动云保存",
            key=f"{key_prefix}_auto_save_enabled",
            help="Only auto-saves after successful generation actions, not on every typing change or rerun.",
        )
        normalize_cloud_widget_keys(key_prefix)
        cloud_name = normalize_project_name_from_cloud_widget(key_prefix)
        force_overwrite = st.checkbox(
            "Force overwrite if cloud changed / 云端变化时仍覆盖",
            value=False,
            key=f"{key_prefix}_force_overwrite",
            help="Default is safer: if cloud changed after your last save, the app blocks silent overwrite.",
        )
        if st.session_state.get("cloud_last_save_at"):
            st.caption(f"Last cloud save: {st.session_state.cloud_last_save_at}")
        c1, c2 = st.columns(2)
        with c1:
            save_disabled = bool(st.session_state.get("ai_running") or st.session_state.get("cloud_save_running"))
            if st.button("💾 Save to cloud", use_container_width=True, key=f"{key_prefix}_save_btn", disabled=save_disabled):
                try:
                    with st.spinner("Saving to cloud..."):
                        saved = cloud_save_project_guarded(username, cloud_name, force=force_overwrite, save_reason="manual save")
                    updated = saved.get("updated_at", "") if isinstance(saved, dict) else ""
                    stamp = updated[:19].replace("T", " ") if updated else st.session_state.get("cloud_last_save_at", "")
                    st.session_state.cloud_last_summary = f"Saved '{cloud_name}' to cloud." + (f" Updated: {stamp}" if stamp else "")
                    st.success(st.session_state.cloud_last_summary)
                except Exception as exc:
                    st.error(f"Cloud save failed: {exc}")
            if save_disabled:
                st.caption("Save is locked while AI is generating or another save is running.")
        with c2:
            if st.button("🔄 Refresh list", use_container_width=True, key=f"{key_prefix}_refresh_btn"):
                st.session_state["_cloud_project_list_nonce"] = time.time()
                st.rerun()

        try:
            projects = cloud_list_projects(username)
        except Exception as exc:
            st.error(f"Could not list cloud projects: {exc}")
            projects = []

        if st.session_state.get("cloud_last_summary"):
            st.caption(st.session_state.cloud_last_summary)
        if st.session_state.get("cloud_snapshot_status"):
            st.caption("Snapshot: " + str(st.session_state.cloud_snapshot_status))

        if not projects:
            st.info("No cloud projects yet for this login. Click Save to cloud first.")
            return

        st.caption(f"Loaded {len(projects)} cloud project record(s) for user: {username}")

        labels = []
        label_to_project: dict[str, dict] = {}
        seen_labels: dict[str, int] = {}
        for item in projects:
            name = item.get("project_name", "Untitled")
            title = item.get("novel_title") or name
            stage = item.get("current_stage", "?")
            updated = item.get("updated_at", "")
            row_id = str(item.get("id", ""))
            short_id = row_id[:8] if row_id else "no-id"
            base_label = f"{name} | 《{title}》 | Stage {stage} | {updated[:19].replace('T', ' ')} | id:{short_id}"
            duplicate_count = seen_labels.get(base_label, 0)
            seen_labels[base_label] = duplicate_count + 1
            label = base_label if duplicate_count == 0 else f"{base_label} | duplicate:{duplicate_count + 1}"
            labels.append(label)
            label_to_project[label] = item

        selected = st.selectbox("Saved cloud projects", labels, key=f"{key_prefix}_saved_project_select")
        b1, b2 = st.columns(2)
        with b1:
            if st.button("📂 Load selected", use_container_width=True, key=f"{key_prefix}_load_btn"):
                try:
                    project_name = str(label_to_project[selected].get("project_name", ""))
                    with st.spinner("Loading from cloud..."):
                        data = cloud_load_project(username, project_name)
                        load_project(data)
                        st.session_state.last_checkpoint_summary = loaded_project_summary()
                    st.success(st.session_state.last_checkpoint_summary)
                    st.rerun()
                except Exception as exc:
                    st.error(f"Cloud load failed: {exc}")
        with b2:
            confirm_delete = st.checkbox("Delete mode", key=f"{key_prefix}_delete_mode")
            if confirm_delete and st.button("🗑️ Delete selected", use_container_width=True, key=f"{key_prefix}_delete_btn"):
                try:
                    project_name = str(label_to_project[selected].get("project_name", ""))
                    cloud_delete_project(username, project_name)
                    st.success(f"Deleted cloud project: {project_name}")
                    st.rerun()
                except Exception as exc:
                    st.error(f"Cloud delete failed: {exc}")

        selected_project_name = str(label_to_project[selected].get("project_name", ""))
        with st.expander("🕘 History / Restore cloud snapshots", expanded=False):
            st.caption("Shows saved snapshots from novel_project_snapshots. If this table is missing, create it in Supabase first.")
            try:
                snapshots = cloud_list_project_snapshots(username, selected_project_name, limit=30)
            except Exception as exc:
                st.warning(f"Could not list snapshots: {exc}")
                st.code(
                    """create table if not exists novel_project_snapshots (
  id uuid primary key default gen_random_uuid(),
  username text not null,
  project_name text not null,
  novel_title text,
  current_stage int,
  save_reason text,
  project_hash text,
  project_data jsonb not null,
  saved_at timestamptz default now()
);

create index if not exists idx_snapshots_user_project_time
on novel_project_snapshots (username, project_name, saved_at desc);""",
                    language="sql",
                )
                snapshots = []
            if not snapshots:
                st.info("No snapshots yet for this project. Save to cloud once after creating the snapshots table.")
            else:
                labels2 = []
                label_to_snapshot = {}
                for snap in snapshots:
                    saved_at = str(snap.get("saved_at", ""))[:19].replace("T", " ")
                    stage = snap.get("current_stage", "?")
                    reason = snap.get("save_reason") or "save"
                    label = f"{saved_at} | Stage {stage} | {reason}"
                    labels2.append(label)
                    label_to_snapshot[label] = snap
                selected_snapshot_label = st.selectbox("Saved snapshots", labels2, key=f"{key_prefix}_snapshot_select")
                snap = label_to_snapshot[selected_snapshot_label]
                st.caption(f"Hash: {snap.get('project_hash', '')}")
                restore_col, copy_col = st.columns(2)
                with restore_col:
                    if st.button("Restore snapshot to current session", use_container_width=True, key=f"{key_prefix}_restore_snapshot"):
                        try:
                            data = cloud_load_project_snapshot(str(snap.get("id")))
                            load_project(data)
                            st.session_state.last_checkpoint_summary = loaded_project_summary() + " Restored from snapshot; Save to cloud to make it latest."
                            st.success(st.session_state.last_checkpoint_summary)
                            st.rerun()
                        except Exception as exc:
                            st.error(f"Snapshot restore failed: {exc}")
                with copy_col:
                    st.caption("After restore, use Save as new copy if you want to keep both current and restored versions.")


def safe_filename(name: str) -> str:
    cleaned = re.sub(r"[^A-Za-z0-9_\-\u4e00-\u9fff]+", "_", name).strip("_")
    return cleaned or "novel_project"


def rebuild_full_draft() -> None:
    chapters = st.session_state.chapters or []
    drafts = st.session_state.chapter_drafts or {}
    parts = []
    for i, ch in enumerate(chapters):
        if i in drafts and drafts[i]:
            heading = chapter_heading(ch)
            parts.append(f"## {heading}\n\n{drafts[i]}")
    st.session_state.full_draft = "\n\n---\n\n".join(parts)


def chapter_heading(ch: dict) -> str:
    num = ch.get("num", "")
    title = ch.get("title", "")
    if "Chinese" in st.session_state.language or "Bilingual" in st.session_state.language:
        return f"第{num}章 {title}"
    return f"Chapter {num}: {title}"


def chinese_num_to_int(raw: str, fallback: int) -> int:
    try:
        return int(raw)
    except Exception:
        pass
    digits = {"零": 0, "一": 1, "二": 2, "两": 2, "三": 3, "四": 4, "五": 5, "六": 6, "七": 7, "八": 8, "九": 9}
    if raw == "十":
        return 10
    if "十" in raw:
        left, _, right = raw.partition("十")
        tens = digits.get(left, 1) if left else 1
        ones = digits.get(right, 0) if right else 0
        return tens * 10 + ones
    return digits.get(raw, fallback)


def _clean_chapter_title(title: str) -> str:
    """Clean markdown/translation noise around a chapter title."""
    title = re.sub(r"^[#>*\-\s]+", "", str(title)).strip()
    title = re.sub(r"\*\*|__|`", "", title).strip()
    title = re.sub(r"\s+", " ", title).strip(" ：:-—–")
    return title or "未命名章节"


def _append_chapter(chapters: list[dict], seen_nums: set[int], num: int, title: str, summary: str = "") -> None:
    if num in seen_nums:
        return
    title = _clean_chapter_title(title)
    summary = str(summary or "").strip().lstrip("—–-:： ").strip()
    chapters.append({"num": num, "title": title, "summary": summary})
    seen_nums.add(num)


def parse_chapters(outline_text: str) -> list[dict]:
    """Robustly parse chapter headings from AI outlines.

    Supports common Chinese/English outline styles, with or without summaries:
    - Chapter 1: Title — summary
    - Chapter 1 - Title
    - 第1章：标题 — 摘要
    - ## 第四章 标题
    - **第4章：标题**
    - 1. 标题 / 一、标题

    The old parser required a title-summary dash, so it could miss most chapters.
    This parser keeps the chapter even when the summary is empty.
    """
    chapters: list[dict] = []
    seen_nums: set[int] = set()
    lines = outline_text.splitlines()

    heading_patterns = [
        # Chapter 1: Title — summary OR Chapter 1 - Title
        r"^Chapter\s+([0-9]+)\s*[:：\-.、]?\s*(.+?)(?:\s*[—–-]\s*(.+))?$",
        # Ch.1: Title
        r"^Ch\.?\s*([0-9]+)\s*[:：\-.、]?\s*(.+?)(?:\s*[—–-]\s*(.+))?$",
        # 第1章：标题 — 摘要 / 第四章 标题
        r"^第\s*([0-9一二两三四五六七八九十百]+)\s*[章节章回卷部]+\s*[:：\-.、]?\s*(.+?)(?:\s*[—–-]\s*(.+))?$",
        # 1. 标题 / 1、标题 / 1) 标题
        r"^([0-9]{1,3})\s*[\.、)）]\s*(.+?)(?:\s*[—–-]\s*(.+))?$",
        # 一、标题 / 四、标题
        r"^([一二两三四五六七八九十百]+)\s*[、\.．)]\s*(.+?)(?:\s*[—–-]\s*(.+))?$",
    ]

    skip_foreshadowing_section = False
    for raw_line in lines:
        raw_clean = raw_line.strip()
        is_markdown_heading = bool(re.match(r"^\s{0,4}#{1,6}\s+", raw_clean))
        clean = raw_clean
        if not clean:
            continue
        clean = re.sub(r"^\s{0,4}(#{1,6})\s*", "", clean).strip()
        heading_plain = re.sub(r"\*\*|__|`", "", clean).strip()
        if is_markdown_heading and re.match(r"^(Foreshadowing|Foreshadowing\s+Map|伏笔|线索|线索回收)", heading_plain, re.I):
            skip_foreshadowing_section = True
            continue
        if is_markdown_heading and re.match(r"^(Chapter\s+Outline|章节大纲)", heading_plain, re.I):
            skip_foreshadowing_section = False
            continue
        if skip_foreshadowing_section:
            continue
        clean = re.sub(r"^[-*+]\s+", "", clean).strip()
        clean = re.sub(r"\*\*|__|`", "", clean).strip()
        clean = clean.strip(" ")
        if not clean:
            continue

        # Avoid parsing section headings such as "## Chapter Outline" as chapters.
        if re.match(r"^(Chapter\s+Outline|章节大纲|大纲|Story\s+Arc|Foreshadowing|Characters)\b", clean, re.I):
            continue

        for pat in heading_patterns:
            m = re.match(pat, clean, re.IGNORECASE)
            if not m:
                continue
            raw_num, title = m.group(1), m.group(2)
            summary = m.group(3) if len(m.groups()) >= 3 else ""
            num = chinese_num_to_int(raw_num, len(chapters) + 1)
            # Avoid treating very long paragraph lines as titles.
            if len(title) > 160 and not summary:
                title, summary = title[:60], title[60:]
            _append_chapter(chapters, seen_nums, num, title, summary)
            break

    # If no formal chapter headings exist, try to infer from lines containing chapter words.
    if not chapters:
        for raw_line in lines:
            clean = raw_line.strip().lstrip("#-*+ >").strip()
            clean = re.sub(r"\*\*|__|`", "", clean).strip()
            if re.search(r"chapter\s+\d+|第\s*[0-9一二两三四五六七八九十百]+\s*[章节章回]", clean, re.IGNORECASE):
                _append_chapter(chapters, seen_nums, len(chapters) + 1, clean, "")

    chapters.sort(key=lambda c: c.get("num", 0))
    return chapters


def infer_chapters_from_story_bible_timeline(story_bible: str) -> list[dict]:
    """Infer provisional chapters from the Story Bible timeline section.

    This helps resume Stage 3 checkpoints that contain a Story Bible but no
    generated outline yet. It is intentionally conservative: it only looks at
    lines under a timeline heading and creates editable placeholder chapters.
    """
    if not story_bible or not story_bible.strip():
        return []

    lines = story_bible.splitlines()
    in_timeline = False
    timeline_lines: list[str] = []

    for raw in lines:
        clean = raw.strip()
        heading = re.sub(r"^#+\s*", "", clean).strip()
        heading_plain = re.sub(r"\*\*|__|`", "", heading).strip()

        if re.search(r"时间线|Timeline", heading_plain, re.IGNORECASE):
            in_timeline = True
            continue

        if in_timeline and clean.startswith("##") and not re.search(r"时间线|Timeline", heading_plain, re.IGNORECASE):
            break

        if in_timeline and clean:
            timeline_lines.append(clean)

    if not timeline_lines:
        return []

    chapters: list[dict] = []
    seen_nums: set[int] = set()
    for raw in timeline_lines:
        clean = raw.strip().lstrip("-*+ >").strip()
        clean = re.sub(r"\*\*|__|`", "", clean).strip()
        if not clean:
            continue
        m = re.match(r"^([0-9一二两三四五六七八九十百]+)\s*[\.、)）]?\s*(.+)$", clean)
        if not m:
            continue
        raw_num = m.group(1)
        body = m.group(2).strip()
        num = chinese_num_to_int(raw_num, len(chapters) + 1)

        # Examples:
        # 第一阶段：枯萎期 —— 展现...
        # Phase 1: The Return - ...
        title = body
        summary = ""
        body = re.sub(r"^第?[一二两三四五六七八九十百0-9]+阶段[：:]?", "", body).strip()
        if "——" in body:
            title, summary = body.split("——", 1)
        elif "—" in body:
            title, summary = body.split("—", 1)
        elif " - " in body:
            title, summary = body.split(" - ", 1)
        elif "：" in body:
            maybe_title, maybe_summary = body.split("：", 1)
            if len(maybe_title) <= 30:
                title, summary = maybe_title, maybe_summary

        title = re.sub(r"^[（(].*?[）)]", "", title).strip(" ：:-—")
        title = title or f"阶段 {num}"
        summary = summary.strip() or clean
        _append_chapter(chapters, seen_nums, num, title, summary)

    chapters.sort(key=lambda c: c.get("num", 0))
    return chapters


def build_recovered_outline_from_chapters(title: str, chapters: list[dict]) -> str:
    """Build an editable temporary outline from inferred chapters."""
    parts = [f"# {title}", "", "## Chapter Outline / 临时章节大纲"]
    for ch in chapters:
        num = ch.get("num", len(parts))
        ch_title = ch.get("title", f"阶段 {num}")
        summary = ch.get("summary", "")
        parts.append(f"第{num}章：{ch_title} — {summary}")
    parts.append("")
    parts.append("## 恢复说明")
    parts.append("此大纲由 Story Bible 的时间线临时恢复，用于让项目可以继续进入 Chapter Planner / Draft Writer。建议在 Stage 3 重新生成或手动微调正式大纲。")
    return "\n".join(parts)


def desired_chapter_count() -> int:
    """Infer the target chapter count from the project length selector."""
    txt = str(st.session_state.get("target_length", "") or "")
    m = re.search(r"~(\d+)", txt)
    if m:
        return int(m.group(1))
    m = re.search(r"(\d+)\s*chapters", txt, re.I)
    if m:
        return int(m.group(1))
    return max(5, len(st.session_state.get("chapters", []) or []), 10)


def build_local_outline_from_material() -> str:
    """Create a usable outline without calling AI. This is a safe fallback for resumed projects."""
    title = st.session_state.get("novel_title", "Untitled") or "Untitled"
    existing = st.session_state.get("chapters", []) or parse_chapters(st.session_state.get("outline", "") or "")
    if not existing and st.session_state.get("story_bible"):
        existing = infer_chapters_from_story_bible_timeline(st.session_state.story_bible)
    if not existing:
        n = desired_chapter_count()
        existing = [{"num": i, "title": f"第{i}阶段", "summary": "请根据故事设定补充本章目标、冲突与钩子。"} for i in range(1, n + 1)]

    parts = [f"# {title}", "", "## Characters", "请参考 Story Bible 中的主要角色设定。", "", "## Story Arc", "请参考 Story Bible 的核心前提、人物关系、时间线与重要伏笔。", "", "## Chapter Outline"]
    for i, ch in enumerate(existing, 1):
        num = ch.get("num", i)
        ch_title = ch.get("title", f"阶段 {num}")
        summary = ch.get("summary", "") or "围绕本阶段目标推进冲突，并在结尾留下下一章钩子。"
        parts.append(f"第{num}章：{ch_title} — {summary}")
    parts.extend(["", "## Foreshadowing Map", "请在后续正式大纲中补充关键线索、回收章节与情绪/商业钩子。"] )
    return "\n".join(parts)


def ensure_chapters_available() -> int:
    """Repair chapters from outline or Story Bible and return count."""
    chapters = st.session_state.get("chapters", []) or []
    if chapters:
        return len(chapters)
    if st.session_state.get("outline"):
        chapters = parse_chapters(st.session_state.outline)
    if not chapters and st.session_state.get("story_bible"):
        chapters = infer_chapters_from_story_bible_timeline(st.session_state.story_bible)
        if chapters and not st.session_state.get("outline"):
            st.session_state.outline = build_recovered_outline_from_chapters(st.session_state.get("novel_title", "Untitled"), chapters)
    st.session_state.chapters = chapters or []
    return len(st.session_state.chapters)


def set_last_action(msg: str) -> None:
    st.session_state["_last_action"] = msg


def compact_story_bible(limit: int = 3500) -> str:
    parts = [
        ("Story Bible", st.session_state.story_bible),
        ("Characters", st.session_state.characters),
        ("Worldbuilding", st.session_state.worldbuilding),
        ("Timeline", st.session_state.timeline),
        ("Glossary / Terms", st.session_state.glossary),
        ("Banned phrases", st.session_state.banned_phrases),
    ]
    text = "\n\n".join(f"## {label}\n{value}" for label, value in parts if str(value).strip())
    return text[:limit]



def platform_strategy_text() -> str:
    """Build platform-positioning guidance for all generation/review prompts."""
    platforms = st.session_state.get("publishing_platforms", []) or []
    submission_types = st.session_state.get("submission_types", []) or []
    goals = st.session_state.get("platform_goals", []) or []
    custom = st.session_state.get("custom_platform_note", "").strip()
    if not platforms and not submission_types and not goals and not custom:
        return ""

    strategy_map = {
        "起点中文网": "起点方向：重视世界观硬度、升级体系、主线目标、资源争夺、长线副本和男频追读节奏。前三章要立住目标和金手指/核心矛盾。",
        "晋江文学城": "晋江方向：重视人设稳定、人物关系、情绪张力、CP/强强互动、文笔质感和角色自主性。冲突要有情感后果。",
        "番茄小说": "番茄方向：重视开局速度、冲突密度、爽点频率、易读性、反转和章节钩子。少铺垫，先给危机、反击和情绪价值。",
        "七猫小说": "七猫方向：重视强情节、强刺激、清晰矛盾、低阅读门槛和高频反转，避免复杂设定拖慢节奏。",
        "掌阅": "掌阅方向：重视成熟商业节奏、强情绪价值、清晰人设和稳定更新感，章节推进要直接。",
        "红袖添香": "红袖方向：重视女频情绪、家族/宫斗/宅斗/复仇线、女性成长和感情张力。",
        "潇湘书院": "潇湘方向：重视大女主、重生复仇、成长线、家族/权谋/事业线与感情线的平衡。",
        "知乎盐选 / 短篇付费": "知乎盐选方向：重视强钩子、短篇完成度、现实/悬疑/复仇反转、信息差和结尾回收。开头必须迅速提出问题。",
        "飞卢": "飞卢方向：重视脑洞直给、系统/金手指、高频爽点、节奏极快和明确升级反馈。",
        "刺猬猫 / 轻小说": "刺猬猫/轻小说方向：重视人设萌点、设定趣味、吐槽感、二次元读者接受度和角色关系张力。",
        "自定义/多平台": "多平台方向：先保证开局钩子、人物卖点、核心爽点和设定清晰，再根据主平台微调节奏。",
    }
    lines = []
    if platforms:
        lines.append("目标投稿平台: " + ", ".join(platforms))
        lines.extend(strategy_map[p] for p in platforms if p in strategy_map)
    if submission_types:
        lines.append("投稿/文章类型: " + ", ".join(submission_types))
    if goals:
        lines.append("平台目标: " + ", ".join(goals))
    if custom:
        lines.append("自定义平台要求: " + custom)
    return "\n".join(lines)


def default_platform_for_review() -> str:
    platforms = st.session_state.get("publishing_platforms", []) or []
    if not platforms:
        return "自动推荐"
    first = platforms[0]
    mapping = {
        "起点中文网": "起点",
        "晋江文学城": "晋江",
        "番茄小说": "番茄",
        "七猫小说": "七猫/掌阅",
        "掌阅": "七猫/掌阅",
        "红袖添香": "红袖/潇湘",
        "潇湘书院": "红袖/潇湘",
        "知乎盐选 / 短篇付费": "知乎盐选",
        "飞卢": "自定义/多平台",
        "刺猬猫 / 轻小说": "自定义/多平台",
        "自定义/多平台": "自定义/多平台",
    }
    return mapping.get(first, "自动推荐")


def default_judges_for_project() -> list[str]:
    platforms = st.session_state.get("publishing_platforms", []) or []
    channel = st.session_state.get("audience_channel", "女频")
    judges: list[str] = []
    if "起点中文网" in platforms or channel == "男频":
        judges.append("起点男频读者")
    if "晋江文学城" in platforms or channel == "女频":
        judges.append("晋江女频读者")
    if "番茄小说" in platforms:
        judges.append("番茄免费文读者")
    if "七猫小说" in platforms or "掌阅" in platforms:
        judges.append("七猫/掌阅读者")
    if "红袖添香" in platforms or "潇湘书院" in platforms:
        judges.append("红袖/潇湘女频读者")
    if "知乎盐选 / 短篇付费" in platforms:
        judges.append("知乎盐选读者")
    for j in ["主编/责编", "作者续写风险评委"]:
        if j not in judges:
            judges.append(j)
    # Keep UI readable.
    return judges[:5]

def writing_context() -> str:
    rules = [
        f"Project title: {st.session_state.novel_title}",
        f"Writing language: {st.session_state.language}",
        f"Audience channel: {st.session_state.audience_channel}",
        f"Primary genre: {st.session_state.genre}",
        f"Genre tags: {', '.join(st.session_state.genre_tags or [])}",
        f"Custom genre tags: {st.session_state.custom_genre_tags.strip() or 'None'}",
        f"Chinese style: {st.session_state.chinese_style}",
    ]
    platform_strategy = platform_strategy_text()
    if platform_strategy:
        rules.append("Publishing/platform strategy:\n" + platform_strategy)
    if "Chinese" in st.session_state.language or "Bilingual" in st.session_state.language:
        rules.extend([
            "Write in natural Chinese novel prose unless the user explicitly asks otherwise.",
            "Avoid translation-like sentence structure and stiff exposition.",
            "Use scene action, dialogue, and subtext instead of direct explanation whenever possible.",
            "For web-novel styles, keep conflict clear, pacing strong, and chapter endings hooky.",
            "Follow 爽文 logic when suitable: clear desire, clear obstacle, visible payoff, reversal, and a curiosity hook at the end.",
            "Do not damage existing plot continuity; preserve established clues, powers, legal evidence, relationship facts, and timeline details.",
        ])
        if st.session_state.audience_channel == "男频":
            rules.append("男频 orientation: stronger external conflict, progression, power/resource gains, tactical reversals, and payoff rhythm.")
        elif st.session_state.audience_channel == "女频":
            rules.append("女频 orientation: stronger emotional stakes, relationship tension, inner conflict, identity pressure, and character agency.")
    else:
        rules.append("Write in polished, natural English fiction prose.")
    if st.session_state.style_note.strip():
        rules.append(f"Style note: {st.session_state.style_note.strip()}")
    if st.session_state.global_instruction.strip():
        rules.append(f"Global instruction: {st.session_state.global_instruction.strip()}")
    if st.session_state.banned_phrases.strip():
        rules.append("Avoid these banned or overused phrases when possible:\n" + st.session_state.banned_phrases.strip())
    if compact_story_bible().strip():
        rules.append("Follow this project Story Bible and do not contradict it:\n" + compact_story_bible())
    return "\n".join(f"- {rule}" for rule in rules)


def push_version(chapter_idx: int, text: str, label: str) -> None:
    versions = st.session_state.chapter_versions.setdefault(chapter_idx, [])
    versions.append({"time": datetime.now().isoformat(timespec="seconds"), "label": label, "text": text})


def revision_prompt_for_mode(mode: str, custom: str) -> str:
    presets = {
        "Make it more suspenseful": "Increase suspense, reduce early explanation, plant questions, and make the scene ending stronger.",
        "Make dialogue more natural": "Rewrite dialogue to sound more natural, character-specific, and less formal. Add subtext.",
        "Add Chinese web-novel pacing": "Make pacing clearer and more addictive: stronger conflict, faster scene movement, and a hook ending.",
        "Make prose less AI-like": "Remove generic AI phrasing, repetitive rhythm, over-explanation, and cliché emotional descriptions.",
        "Increase emotional tension": "Increase emotional stakes through action, silence, subtext, and character choices rather than direct explanation.",
        "Add sensory detail": "Add concrete sensory details tied to character perspective without slowing the scene too much.",
        "Reduce exposition": "Cut exposition and replace it with action, dialogue, conflict, and implication.",
        "Add conflict": "Add sharper conflict, resistance, obstacles, and consequences within the scene.",
        "Make ending stronger": "Rewrite the ending to create curiosity, danger, emotional impact, or an unanswered question.",
        "Rewrite in 古风 tone": "Rewrite with a more elegant 古风 tone, but avoid over-decorated language and keep readability.",
    }
    base = presets.get(mode, "")
    if custom.strip():
        return base + "\nAdditional user instruction: " + custom.strip()
    return base

def render_quick_reference(chapter_idx: int | None = None) -> None:
    """Always-available setting lookup widget to protect flow while drafting/editing."""
    title = "📌 设定速查小挂件 / Story Bible Quick Reference"
    with st.expander(title, expanded=False):
        meta = [
            f"**书名**：{st.session_state.novel_title}",
            f"**频道**：{st.session_state.audience_channel}",
            f"**主类型**：{st.session_state.genre}",
            f"**多题材标签**：{', '.join(st.session_state.genre_tags or []) or '未选择'}",
            f"**自定义标签**：{st.session_state.custom_genre_tags or '无'}",
            f"**风格**：{st.session_state.chinese_style}",
        ]
        st.markdown("  \n".join(meta))
        if chapter_idx is not None and chapter_idx < len(st.session_state.chapters):
            ch = st.session_state.chapters[chapter_idx]
            st.markdown(f"**当前章节**：{chapter_heading(ch)}")
            st.markdown(f"**章节摘要**：{ch.get('summary', '') or '无'}")
            beat = st.session_state.chapter_beats.get(chapter_idx, "")
            if beat:
                st.markdown("**本章 Beat Sheet**")
                st.text_area("Beat Sheet", value=beat, height=180, key=f"quick_beat_{chapter_idx}", label_visibility="collapsed")
        tabs = st.tabs(["人物", "世界规则", "时间线", "术语/伏笔", "禁用套话", "总设定"])
        with tabs[0]:
            st.text_area("Characters", value=st.session_state.characters or "暂无人物设定", height=180, key=f"quick_char_{chapter_idx}", label_visibility="collapsed")
        with tabs[1]:
            st.text_area("Worldbuilding", value=st.session_state.worldbuilding or "暂无世界观/规则", height=180, key=f"quick_world_{chapter_idx}", label_visibility="collapsed")
        with tabs[2]:
            st.text_area("Timeline", value=st.session_state.timeline or "暂无时间线", height=180, key=f"quick_timeline_{chapter_idx}", label_visibility="collapsed")
        with tabs[3]:
            st.text_area("Glossary", value=st.session_state.glossary or "暂无术语/伏笔", height=180, key=f"quick_glossary_{chapter_idx}", label_visibility="collapsed")
        with tabs[4]:
            st.text_area("Banned phrases", value=st.session_state.banned_phrases or "暂无禁用套话", height=180, key=f"quick_banned_{chapter_idx}", label_visibility="collapsed")
        with tabs[5]:
            st.text_area("Story Bible", value=st.session_state.story_bible or "暂无总设定", height=220, key=f"quick_bible_{chapter_idx}", label_visibility="collapsed")


def _chapter_generation_user_prompt(chapter_idx: int, length_rule: str) -> str:
    ch = st.session_state.chapters[chapter_idx]
    beat = st.session_state.chapter_beats.get(chapter_idx, "")
    previous_context = ""
    if chapter_idx > 0 and st.session_state.chapter_drafts.get(chapter_idx - 1):
        previous_context = st.session_state.chapter_drafts[chapter_idx - 1][-1600:]
    return f"""Writing context:
{writing_context()}

Overall outline:
{st.session_state.outline[:5000]}

Previous chapter ending, if available:
{previous_context or 'None'}

Chapter:
{chapter_heading(ch)}
Summary: {ch.get('summary', '')}

Beat sheet:
{beat or 'No beat sheet provided. Use the outline summary.'}

Length target: {length_rule}

Draft the chapter as polished prose.
Requirements:
- Follow the Story Bible and current chapter beat sheet.
- Keep 爽文 logic when appropriate: conflict, reversal, payoff, and hook.
- Use natural Chinese/English rhythm based on the selected language.
- Avoid robotic phrasing and banned phrases.
- Do not include planning notes, outline bullets, checks, self-correction, analysis, or prompt text.
- Start directly with the finished chapter prose. Return only the chapter text.
"""


def generate_chapter_text(chapter_idx: int, length_rule: str, placeholder) -> str:
    return generate_text(
        system="You are a skilled fiction writer. Draft vivid, natural, scene-driven prose while preserving continuity.",
        user=_chapter_generation_user_prompt(chapter_idx, length_rule),
        placeholder=placeholder,
        max_tokens=5500,
        cleaner=clean_chapter_draft_output,
    )


def render_generate_full_text_panel(location_key: str, default_length_label: str = "Normal chapter") -> None:
    """One-click full manuscript generation / missing chapter fill for stage 5 and later."""
    if not st.session_state.chapters:
        return
    with st.expander("🚀 一键生成全文 / 补齐未写章节", expanded=False):
        st.caption("建议先生成 Story Bible 和 Chapter Beat Sheet。按钮会按章节顺序写作；默认只补齐未写章节，避免覆盖已有正文。")
        length_options = ["Short scene", "Normal chapter", "Long chapter"]
        default_index = length_options.index(default_length_label) if default_length_label in length_options else 1
        full_length_target = st.selectbox("全文生成长度档", length_options, index=default_index, key=f"full_len_{location_key}")
        length_rule = {
            "Short scene": "700-1000 English words or 1000-1800 Chinese characters",
            "Normal chapter": "1000-1600 English words or 1800-3000 Chinese characters",
            "Long chapter": "1600-2400 English words or 3000-4500 Chinese characters",
        }[full_length_target]
        overwrite = st.checkbox("覆盖已有章节（危险，默认关闭）", value=False, key=f"full_overwrite_{location_key}")
        max_count = st.number_input("本次最多生成章节数（防止一次调用太久）", min_value=1, max_value=max(1, len(st.session_state.chapters)), value=min(5, len(st.session_state.chapters)), key=f"full_max_{location_key}")
        if st.button("一键生成 / 补齐全文 ✨", type="primary", key=f"full_btn_{location_key}"):
            targets = []
            for i in range(len(st.session_state.chapters)):
                if overwrite or not st.session_state.chapter_drafts.get(i):
                    targets.append(i)
                if len(targets) >= int(max_count):
                    break
            if not targets:
                st.info("没有需要生成的章节。勾选覆盖已有章节，或增加新章节后再试。")
            else:
                progress = st.progress(0)
                status = st.empty()
                for n, chapter_idx in enumerate(targets, 1):
                    ch = st.session_state.chapters[chapter_idx]
                    status.info(f"正在生成 {chapter_heading(ch)} ({n}/{len(targets)})...")
                    placeholder = st.empty()
                    result = generate_chapter_text(chapter_idx, length_rule, placeholder)
                    if result:
                        if st.session_state.chapter_drafts.get(chapter_idx):
                            push_version(chapter_idx, st.session_state.chapter_drafts[chapter_idx], "Before one-click full generation")
                        st.session_state.chapter_drafts[chapter_idx] = result
                        rebuild_full_draft()
                        mark_dirty()
                        cloud_auto_save(f"one-click draft {chapter_heading(ch)}")
                    progress.progress(n / len(targets))
                status.success("本轮全文生成/补齐完成。可再次点击继续生成剩余章节。")
                st.rerun()


def render_local_refinement_panel(chapter_idx: int, draft: str, location_key: str) -> None:
    """Point-kill paragraph refinement that never overwrites the whole chapter accidentally."""
    st.subheader("🎯 局部段落点杀式精修")
    st.caption("粘贴你要精修的一小段。AI 只处理这一段；只有当原文能精确匹配时，才允许安全替换。")
    selected = st.text_area("要精修的原文片段", height=140, key=f"local_src_{location_key}")
    local_mode = st.selectbox(
        "精修目标",
        ["去AI味但保留剧情", "更有爽文节奏", "对白更自然", "增强钩子", "增强情绪张力", "减少解释改成动作", "古风但不堆辞藻", "自定义"],
        key=f"local_mode_{location_key}",
    )
    local_extra = st.text_area("补充要求", height=80, key=f"local_extra_{location_key}", placeholder="例如：保留法律诉讼线索；异能规则不能改；只改语气和节奏。")
    if st.button("生成局部精修稿 ✨", key=f"local_btn_{location_key}"):
        if not selected.strip():
            st.error("请先粘贴要精修的段落。")
        else:
            with st.spinner("局部精修中..."):
                placeholder = st.empty()
                result = generate_text(
                    system="You are a careful line editor. Rewrite only the selected passage and preserve all plot facts.",
                    user=f"""Writing context:
{writing_context()}

Current chapter heading:
{chapter_heading(st.session_state.chapters[chapter_idx]) if chapter_idx < len(st.session_state.chapters) else chapter_idx}

Nearby chapter context:
{draft[:1200]}
...
{draft[-1200:]}

Selected passage to refine:
{selected}

Refinement goal: {local_mode}
Extra instruction: {local_extra or 'None'}

Return only the refined passage. Preserve names, clues, legal evidence, ability rules, timeline facts, and plot meaning.
""",
                    placeholder=placeholder,
                    max_tokens=2500,
                )
            if result:
                st.session_state.local_refinements[chapter_idx] = {"source": selected, "result": result, "time": datetime.now().isoformat(timespec="seconds")}
                cloud_auto_save("local refinement")
                st.success("局部精修稿已生成。先对比，再决定是否安全替换。")
                st.rerun()
    refinement = st.session_state.local_refinements.get(chapter_idx)
    if refinement:
        left, right = st.columns(2)
        with left:
            st.markdown("**原片段**")
            st.text_area("Original selected passage", value=refinement.get("source", ""), height=220, key=f"local_orig_{location_key}")
        with right:
            st.markdown("**精修稿**")
            st.text_area("Refined passage", value=refinement.get("result", ""), height=220, key=f"local_refined_{location_key}")
        a, b, c = st.columns(3)
        with a:
            if st.button("安全替换这一个片段", type="primary", key=f"local_apply_{location_key}"):
                src = refinement.get("source", "")
                res = refinement.get("result", "")
                current = st.session_state.chapter_drafts.get(chapter_idx, "")
                if src and src in current:
                    push_version(chapter_idx, current, "Before local paragraph replacement")
                    st.session_state.chapter_drafts[chapter_idx] = current.replace(src, res, 1)
                    rebuild_full_draft()
                    st.session_state.local_refinements.pop(chapter_idx, None)
                    st.success("已只替换匹配到的这一处片段。")
                    st.rerun()
                else:
                    st.error("没有在当前章节中精确找到原片段。为避免误伤全文，未自动替换；请手动复制精修稿。")
        with b:
            if st.button("保存为版本，不替换", key=f"local_save_{location_key}"):
                push_version(chapter_idx, refinement.get("result", ""), "Local refinement proposal")
                st.success("已保存为版本。")
        with c:
            if st.button("丢弃局部精修稿", key=f"local_discard_{location_key}"):
                st.session_state.local_refinements.pop(chapter_idx, None)
                st.rerun()


def render_batch_revision_panel(location_key: str) -> None:
    """Batch full-draft revision proposals, chapter by chapter, without overwriting originals."""
    with st.expander("🔁 全篇批量修改提案 / Batch full-draft revision", expanded=True):
        if not st.session_state.get("chapter_drafts"):
            st.info("当前还没有已保存的章节正文。请先生成至少一章，或点击上面的“一键生成 / 补齐全文”。")
            return
        st.caption("安全模式：逐章生成修改提案，不直接覆盖原文。你可以逐章接受，或一键接受全部提案。")
        mode = st.selectbox(
            "全篇修改目标",
            [
                "全篇去 AI 味但保留剧情",
                "全篇增强爽文节奏",
                "全篇增强女频情绪张力",
                "全篇统一古风语气",
                "全篇减少解释，增加动作和对白",
                "全篇增强断章钩子",
                "自定义",
            ],
            key=f"batch_mode_{location_key}",
        )
        extra = st.text_area("全篇修改补充要求", height=90, key=f"batch_extra_{location_key}", placeholder="例如：保留设定和伏笔；不要改变人物关系；只改节奏和语言。")
        drafted_indices = sorted([i for i in st.session_state.chapter_drafts.keys() if st.session_state.chapter_drafts.get(i)])
        selected = st.multiselect(
            "选择要生成提案的章节",
            drafted_indices,
            default=drafted_indices[: min(5, len(drafted_indices))],
            format_func=lambda i: chapter_heading(st.session_state.chapters[i]) if i < len(st.session_state.chapters) else f"Chapter {i+1}",
            key=f"batch_selected_{location_key}",
        )
        overwrite_existing = st.checkbox("重新生成已有提案", value=False, key=f"batch_overwrite_{location_key}")
        if st.button("生成全篇逐章修改提案 ✨", type="primary", use_container_width=True, key=f"batch_generate_{location_key}"):
            if not selected:
                st.error("请至少选择一章。")
            else:
                progress = st.progress(0)
                status = st.empty()
                for n, idx in enumerate(selected, 1):
                    if (not overwrite_existing) and st.session_state.batch_revision_proposals.get(idx):
                        progress.progress(n / len(selected))
                        continue
                    ch = st.session_state.chapters[idx] if idx < len(st.session_state.chapters) else {"num": idx+1, "title": f"Chapter {idx+1}", "summary": ""}
                    draft = st.session_state.chapter_drafts.get(idx, "")
                    status.info(f"正在生成提案 {n}/{len(selected)}：{chapter_heading(ch) if idx < len(st.session_state.chapters) else idx+1}")
                    placeholder = st.empty()
                    result = generate_text(
                        system="You are a careful fiction editor. Revise one chapter while preserving plot facts, continuity, names, clues, and chapter meaning. Return only the revised chapter.",
                        user=f"""Writing context:
{writing_context()}

Full outline:
{st.session_state.outline[:5000]}

Chapter:
{chapter_heading(ch) if idx < len(st.session_state.chapters) else f'Chapter {idx+1}'}

Batch revision goal: {mode}
Extra instruction: {extra or 'None'}

Original chapter draft:
{draft}

Return only the revised chapter text. Do not include notes, analysis, or headings unless the original chapter already has a heading.
""",
                        placeholder=placeholder,
                        max_tokens=6500,
                        cleaner=clean_chapter_draft_output,
                        task_name=f"Batch revision {idx+1}",
                    )
                    if result:
                        st.session_state.batch_revision_proposals[idx] = {
                            "time": datetime.now().isoformat(timespec="seconds"),
                            "mode": mode,
                            "extra": extra,
                            "original": draft,
                            "proposal": result,
                        }
                        mark_dirty()
                    progress.progress(n / len(selected))
                status.success("批量修改提案已生成。原文未被覆盖。")
                cloud_auto_save("batch revision proposals")
                st.rerun()

        proposals = st.session_state.get("batch_revision_proposals", {}) or {}
        if proposals:
            st.markdown("### 已生成的全篇修改提案")
            a, b = st.columns(2)
            with a:
                if st.button("接受全部提案", type="primary", use_container_width=True, key=f"batch_accept_all_{location_key}"):
                    for idx, item in list(proposals.items()):
                        current = st.session_state.chapter_drafts.get(idx, "")
                        push_version(idx, current, "Before accepted batch revision")
                        st.session_state.chapter_drafts[idx] = item.get("proposal", "")
                    st.session_state.batch_revision_proposals = {}
                    rebuild_full_draft()
                    mark_dirty()
                    cloud_auto_save("accepted all batch revisions")
                    st.success("已接受全部提案。")
                    st.rerun()
            with b:
                if st.button("清空全部提案", use_container_width=True, key=f"batch_clear_all_{location_key}"):
                    st.session_state.batch_revision_proposals = {}
                    mark_dirty()
                    st.rerun()
            for idx in sorted(proposals.keys()):
                item = proposals[idx]
                label = chapter_heading(st.session_state.chapters[idx]) if idx < len(st.session_state.chapters) else f"Chapter {idx+1}"
                with st.expander(f"{label} — {item.get('mode', '')}"):
                    left, right = st.columns(2)
                    with left:
                        st.markdown("**原文**")
                        st.text_area("Original", value=item.get("original", ""), height=260, key=f"batch_orig_{location_key}_{idx}")
                    with right:
                        st.markdown("**修改提案**")
                        st.text_area("Proposal", value=item.get("proposal", ""), height=260, key=f"batch_prop_{location_key}_{idx}")
                    c1, c2 = st.columns(2)
                    with c1:
                        if st.button("接受这一章", type="primary", key=f"batch_accept_{location_key}_{idx}"):
                            current = st.session_state.chapter_drafts.get(idx, "")
                            push_version(idx, current, "Before accepted batch revision")
                            st.session_state.chapter_drafts[idx] = item.get("proposal", "")
                            st.session_state.batch_revision_proposals.pop(idx, None)
                            rebuild_full_draft()
                            mark_dirty()
                            cloud_auto_save(f"accepted batch revision {label}")
                            st.rerun()
                    with c2:
                        if st.button("放弃这一章提案", key=f"batch_discard_{location_key}_{idx}"):
                            st.session_state.batch_revision_proposals.pop(idx, None)
                            mark_dirty()
                            st.rerun()


def _score_source_text(score_target: str, chapter_idx: int | None = None) -> tuple[str, str]:
    """Return label and text for the market review target."""
    rebuild_full_draft()
    if score_target == "Story Bible / 设定集":
        return "Story Bible / 设定集", compact_story_bible(12000) or st.session_state.story_bible
    if score_target == "Outline / 大纲":
        return "Outline / 大纲", st.session_state.outline
    if score_target == "Current Chapter / 当前章节":
        idx = chapter_idx if chapter_idx is not None else 0
        ch_label = chapter_heading(st.session_state.chapters[idx]) if idx < len(st.session_state.chapters) else f"Chapter {idx + 1}"
        return ch_label, st.session_state.chapter_drafts.get(idx, "") or st.session_state.chapter_beats.get(idx, "")
    return "Full Draft / 全文草稿", st.session_state.full_draft or st.session_state.outline or compact_story_bible(12000)


def _market_review_prompt(score_target: str, source_label: str, source_text: str, judges: list[str], tone: str, target_platform: str, extra: str) -> str:
    return f"""Writing context:
{writing_context()}

评分对象: {score_target}
对象名称: {source_label}
目标平台: {target_platform}
项目投稿定位:
{platform_strategy_text() or '未设置'}
评审语气: {tone}
评委组: {', '.join(judges)}

待评分内容:
{source_text[:12000]}

请模拟一个网文市场评审团，对作品进行商业与创作层面的评估。
只输出最终评分报告，不要展示分析过程、提示词、自检或追问。

必须按以下结构输出中文报告：
## 综合评分
- 综合商业潜力：__/100
- 读者追读潜力：__/100
- 编辑签约潜力：__/100
- 作者可持续写作度：__/100
- 推荐平台：____
- 一句话结论：____

## 分项评分表
请用表格，包含：维度、分数、判断、优先级。
维度至少包括：开局钩子、人设吸引力、爽点密度、情绪张力、题材匹配度、平台适配度、设定清晰度、长期连载潜力、AI味风险、写崩风险。

## 各评委意见
分别模拟所选评委。每个评委必须有：最喜欢什么、最担心什么、会不会继续看、修改建议。

## 最大卖点
列出 3-5 个真正能吸引读者/编辑的卖点。

## 最大风险
列出 3-5 个可能导致弃文、写崩、签约失败或追读下降的风险。

## 前三条必须修改
给出具体、可执行的修改动作，不要泛泛而谈。

## 下一步修改指令
给作者一段可以直接复制到本 app 里使用的修改指令。

额外要求:
{extra or '无'}
"""


def run_market_review(score_target: str, chapter_idx: int | None, judges: list[str], tone: str, target_platform: str, extra: str, key: str) -> None:
    """Generate and store market/web-novel review."""
    source_label, source_text = _score_source_text(score_target, chapter_idx)
    if not source_text.strip():
        st.error("当前评分对象没有内容。请先生成设定集、大纲、章节或全文草稿。")
        return
    if not judges:
        st.error("请至少选择一个评委。")
        return
    with st.spinner("网文评审团正在打分..."):
        placeholder = st.empty()
        result = generate_text(
            system="You are a web-novel market review panel. Be commercially sharp, specific, and actionable. Return only the final Chinese review report.",
            user=_market_review_prompt(score_target, source_label, source_text, judges, tone, target_platform, extra),
            placeholder=placeholder,
            max_tokens=4500,
        )
    if result:
        st.session_state.market_reviews[key] = {
            "time": datetime.now().isoformat(timespec="seconds"),
            "score_target": score_target,
            "source_label": source_label,
            "judges": judges,
            "tone": tone,
            "target_platform": target_platform,
            "extra": extra,
            "report": result,
        }
        cloud_auto_save(f"market review: {source_label}")
        st.success("评分完成。")


def render_market_review_controls(location_key: str, compact: bool = False, chapter_idx: int | None = None, show_inline_report: bool = True) -> None:
    """Reusable review panel for global stage and current chapter quick score.

    show_inline_report=False is used on the dedicated Web Novel Score stage because
    that page already renders saved reports in the Historical scoring reports
    section. Without this, the same newly generated report appears twice: once
    under the controls and once in history.
    """
    default_target = "Current Chapter / 当前章节" if chapter_idx is not None else "Story Bible / 设定集"
    target_options = ["Story Bible / 设定集", "Outline / 大纲", "Current Chapter / 当前章节", "Full Draft / 全文草稿"]
    if compact:
        st.caption("让不同平台读者和主编视角快速判断：能不能追、能不能签、哪里最该改。")
    col1, col2 = st.columns(2)
    with col1:
        score_target = st.selectbox("评分对象", target_options, index=target_options.index(default_target), key=f"score_target_{location_key}")
        platform_choices = ["自动推荐", "起点", "晋江", "番茄", "七猫/掌阅", "红袖/潇湘", "知乎盐选", "自定义/多平台"]
        default_platform = default_platform_for_review()
        target_platform = st.selectbox(
            "目标平台",
            platform_choices,
            index=platform_choices.index(default_platform) if default_platform in platform_choices else 0,
            key=f"score_platform_{location_key}",
        )
    with col2:
        tone = st.selectbox("评审语气", ["专业编辑", "温和鼓励", "毒舌主编"], key=f"score_tone_{location_key}")
        selected_chapter_idx = chapter_idx
        if score_target == "Current Chapter / 当前章节" and chapter_idx is None and st.session_state.chapters:
            chapter_labels = [chapter_heading(ch) for ch in st.session_state.chapters]
            selected_chapter_idx = st.selectbox("选择章节", range(len(chapter_labels)), format_func=lambda i: chapter_labels[i], key=f"score_chapter_{location_key}")
    default_judges = default_judges_for_project()
    judges = st.multiselect(
        "评委组",
        ["起点男频读者", "晋江女频读者", "番茄免费文读者", "七猫/掌阅读者", "红袖/潇湘女频读者", "知乎盐选读者", "主编/责编", "作者续写风险评委"],
        default=default_judges,
        key=f"score_judges_{location_key}",
    )
    extra = st.text_area("额外评分要求", height=80, placeholder="例如：重点看前三章留存、断章钩子、女频情绪张力、AI味风险。", key=f"score_extra_{location_key}")
    report_key = f"{location_key}_{score_target}_{selected_chapter_idx if selected_chapter_idx is not None else 'global'}"
    if st.button("开始评分 📊", type="primary", key=f"score_btn_{location_key}"):
        run_market_review(score_target, selected_chapter_idx, judges, tone, target_platform, extra, report_key)
    review = st.session_state.market_reviews.get(report_key)
    if review and show_inline_report:
        st.markdown(review.get("report", ""))
        st.download_button(
            "Download scoring report Markdown",
            data=review.get("report", ""),
            file_name=f"{safe_filename(st.session_state.project_name)}_market_score.md",
            mime="text/markdown",
            use_container_width=True,
            key=f"score_download_{location_key}",
        )


# =============================================================================
# Multi-user app login
# =============================================================================
def _configured_users() -> dict[str, str]:
    """Read multiple users from Streamlit secrets or environment.

    Preferred Streamlit secrets format:

    [users]
    haopeng = "password1"
    editor = "password2"

    Fallback for simple single-password deployments:
    AUTH_USERNAME = "haopeng"
    AUTH_PASSWORD = "password1"
    """
    users: dict[str, str] = {}

    try:
        raw_users = st.secrets.get("users", {})
        if raw_users:
            users = {
                str(k).strip(): str(v)
                for k, v in dict(raw_users).items()
                if str(k).strip() and str(v)
            }
    except Exception:
        users = {}

    # Optional environment variable format:
    # APP_USERS='{"haopeng":"password1","editor":"password2"}'
    if not users:
        raw_env_users = os.getenv("APP_USERS", "").strip()
        if raw_env_users:
            try:
                parsed = json.loads(raw_env_users)
                if isinstance(parsed, dict):
                    users = {
                        str(k).strip(): str(v)
                        for k, v in parsed.items()
                        if str(k).strip() and str(v)
                    }
            except Exception:
                users = {}

    # Backward-compatible single-user fallback.
    if not users:
        username = _read_secret_or_env("AUTH_USERNAME") or "user"
        password = _read_secret_or_env("AUTH_PASSWORD")
        if password:
            users = {username: password}

    return users


def require_login() -> None:
    """Protect the Streamlit app with username/password login for multiple users."""
    restore_login_from_query_token()
    users = _configured_users()

    if not users:
        st.warning(
            "Login is enabled, but no users are configured. "
            "Add a [users] table in Streamlit secrets before sharing the app."
        )
        with st.expander("Expected secrets format"):
            st.code(
                """[users]
haopeng = "your-password"
editor = "another-password"

GEMINI_API_KEY = "your-gemini-key"
BIGMODEL_API_KEY = "your-bigmodel-key"
""",
                language="toml",
            )
        return

    if "auth_ok" not in st.session_state:
        st.session_state.auth_ok = False
    if "auth_user" not in st.session_state:
        st.session_state.auth_user = ""

    if st.session_state.auth_ok:
        with st.sidebar:
            st.caption(f"🔒 Logged in as {st.session_state.auth_user}")
            if st.button("Log out", use_container_width=True):
                st.session_state.auth_ok = False
                st.session_state.auth_user = ""
                clear_login_query_token()
                st.rerun()
        return

    st.markdown("# 🔒 AI Novel Studio Login")
    st.caption("Enter your username and password to continue.")

    with st.form("login_form"):
        entered_username = st.text_input("Username")
        entered_password = st.text_input("Password", type="password")
        submitted = st.form_submit_button("Log in", use_container_width=True)

    if submitted:
        username = entered_username.strip()
        expected_password = users.get(username)
        if expected_password and entered_password == expected_password:
            st.session_state.auth_ok = True
            st.session_state.auth_user = username
            try:
                st.query_params["auth"] = make_auth_token(username)
            except Exception:
                pass
            st.rerun()
        else:
            st.error("Invalid username or password.")

    st.stop()


require_login()
render_status_bar()


# =============================================================================
# Responsive UI helpers: desktop + mobile writing mode
# =============================================================================
def current_view_mode() -> str:
    mode = st.session_state.get("view_mode", "Auto")
    if mode not in ["Auto", "Desktop", "Mobile"]:
        mode = "Auto"
    return mode


def is_mobile_mode() -> bool:
    # Streamlit cannot reliably know browser width without a custom component.
    # Auto keeps desktop layout by default; users can switch Mobile manually on phones.
    return current_view_mode() == "Mobile"


def ui_height(desktop: int, mobile: int | None = None) -> int:
    return mobile if is_mobile_mode() and mobile is not None else desktop


def mobile_stage_map() -> list[tuple[str, int]]:
    return [
        ("Project", 1),
        ("World", 2),
        ("Outline", 3),
        ("Write", 5),
        ("Edit", 6),
        ("Export", 10),
    ]


def render_mobile_top_nav() -> None:
    if not is_mobile_mode():
        return
    st.markdown('<div class="mobile-card">', unsafe_allow_html=True)
    st.markdown("**Mobile writing mode / 手机写作模式**")
    labels = [x[0] for x in mobile_stage_map()]
    stage_by_label = dict(mobile_stage_map())
    current = next((label for label, stage in mobile_stage_map() if stage == st.session_state.get("stage")), "Write")
    selected = st.radio(
        "Quick navigation",
        labels,
        index=labels.index(current) if current in labels else 3,
        horizontal=True,
        label_visibility="collapsed",
        disabled=bool(st.session_state.get("ai_running")),
        key="mobile_top_nav_radio",
    )
    target_stage = stage_by_label.get(selected, st.session_state.get("stage", 5))
    if target_stage != st.session_state.get("stage"):
        st.session_state.stage = target_stage
        st.rerun()
    st.markdown('</div>', unsafe_allow_html=True)


def render_mobile_dashboard() -> None:
    if not is_mobile_mode():
        return
    title = st.session_state.get("novel_title") or "Untitled"
    project = default_cloud_project_name()
    stage = int(st.session_state.get("stage", 1) or 1)
    drafts = st.session_state.get("chapter_drafts") or {}
    chapters = st.session_state.get("chapters") or []
    last_saved = st.session_state.get("cloud_last_save_at") or "not saved in this session"
    st.markdown('<div class="mobile-card">', unsafe_allow_html=True)
    st.markdown(f'<div class="mobile-title">《{title}》</div>', unsafe_allow_html=True)
    st.markdown(f'<div class="mobile-kpi">Project: {project} · Stage {stage} · Drafted chapters: {len(drafts)}/{len(chapters)}</div>', unsafe_allow_html=True)
    st.markdown(f'<div class="mobile-kpi">Last cloud save: {last_saved}</div>', unsafe_allow_html=True)
    c1, c2, c3 = st.columns(3)
    with c1:
        if st.button("✍️ Write", use_container_width=True, disabled=bool(st.session_state.get("ai_running")), key="mobile_dash_write"):
            st.session_state.stage = 5
            st.rerun()
    with c2:
        if st.button("💾 Save", use_container_width=True, disabled=bool(st.session_state.get("ai_running")), key="mobile_dash_save"):
            try:
                cloud_save_project_guarded(st.session_state.get("auth_user", "user"), project, force=True, save_reason="mobile quick save")
                st.success("Saved to cloud.")
            except Exception as exc:
                st.error(f"Cloud save failed: {exc}")
    with c3:
        if st.button("📥 Export", use_container_width=True, disabled=bool(st.session_state.get("ai_running")), key="mobile_dash_export"):
            st.session_state.stage = 10
            st.rerun()
    st.markdown('</div>', unsafe_allow_html=True)

# =============================================================================
# Sidebar
# =============================================================================
STAGES = [
    ("🧭", "Project Setup"),
    ("📚", "Story Bible"),
    ("📝", "Outline"),
    ("🧩", "Chapter Planner"),
    ("✍️", "Draft Writer"),
    ("🔁", "Revision Studio"),
    ("📊", "Web Novel Score"),
    ("🔍", "Consistency Editor"),
    ("✨", "Polish"),
    ("📥", "Export"),
]

with st.sidebar:
    st.markdown("## 📖 AI Novel Studio")
    st.caption("Gemini first → BigModel/Zhipu fallback")
    st.radio(
        "View mode / 显示模式",
        ["Auto", "Desktop", "Mobile"],
        key="view_mode",
        horizontal=True,
        help="Mobile mode simplifies navigation and writing controls. Auto keeps the full desktop workspace.",
    )
    gemini_key = get_gemini_api_key()
    bigmodel_key = get_bigmodel_api_key()
    if gemini_key and bigmodel_key:
        st.success("AI ready: Gemini + BigModel")
    elif gemini_key:
        st.success("AI ready: Gemini only")
    elif bigmodel_key:
        st.success("AI ready: BigModel only")
    else:
        st.warning("Add API keys in Streamlit secrets")
    with st.expander("🧪 AI connection check", expanded=False):
        st.caption("Use this if Generate buttons seem to do nothing.")
        if st.button("Test AI call", use_container_width=True, key="ai_test_call_v12"):
            try:
                set_last_action("Testing AI provider...")
                with st.spinner("Testing AI provider..."):
                    test = call_ai("You are a test responder.", "Reply only: OK", max_tokens=50)
                st.success(f"AI response: {test[:120]}")
                set_last_action("AI test succeeded.")
            except Exception as exc:
                st.error(f"AI test failed: {exc}")
                set_last_action(f"AI test failed: {exc}")
        if st.session_state.get("_last_action"):
            st.caption("Last action:")
            st.code(st.session_state.get("_last_action"), language="text")
    st.markdown("---")
    for i, (icon, label) in enumerate(STAGES, 1):
        btn_type = "primary" if st.session_state.stage == i else "secondary"
        if st.button(
            f"{icon} {i}. {label}",
            use_container_width=True,
            type=btn_type,
            disabled=bool(st.session_state.get("ai_running")),
            help="AI is running. Wait until the current generation finishes." if st.session_state.get("ai_running") else None,
        ):
            st.session_state.stage = i
            st.rerun()
    st.markdown("---")
    st.caption(f"Project: {st.session_state.project_name}")
    if is_mobile_mode():
        with st.expander("Cloud projects / 云端项目", expanded=False):
            render_cloud_projects_panel("sidebar_cloud")
    else:
        render_cloud_projects_panel("sidebar_cloud")

    # Persistent checkpoint: available from every phase.
    # This lets the author save the exact current state, upload later,
    # and resume from the same stage without losing writing flow.
    current_stage_label = next((label for n, (_, label) in enumerate(STAGES, 1) if n == st.session_state.stage), "Current Stage")
    with st.expander("💾 Save / Load checkpoint", expanded=False):
        checkpoint = project_data()
        checkpoint["checkpoint_stage"] = st.session_state.stage
        checkpoint["checkpoint_stage_label"] = current_stage_label
        checkpoint_name = f"{safe_filename(st.session_state.project_name)}_stage{st.session_state.stage}_{safe_filename(current_stage_label)}.json"
        st.download_button(
            "Download current phase JSON",
            data=json.dumps(checkpoint, ensure_ascii=False, indent=2),
            file_name=checkpoint_name,
            mime="application/json",
            use_container_width=True,
            help="Save everything in the project, including the current phase, chapters, drafts, revisions, and Story Bible.",
        )
        if st.session_state.get("last_checkpoint_summary"):
            st.success(st.session_state.last_checkpoint_summary)
        uploaded_checkpoint = st.file_uploader(
            "Resume from JSON",
            type=["json"],
            key="sidebar_checkpoint_upload",
            help="Upload a project/checkpoint JSON and continue from the saved phase.",
        )
        if uploaded_checkpoint is not None:
            try:
                raw_checkpoint = uploaded_checkpoint.getvalue()
                checkpoint_hash = hashlib.md5(raw_checkpoint).hexdigest()
                last_hash = st.session_state.get("_last_loaded_checkpoint_hash")

                if checkpoint_hash != last_hash:
                    with st.spinner("正在读取项目 JSON / Loading checkpoint..."):
                        data = json.loads(raw_checkpoint.decode("utf-8"))
                        st.session_state["_last_loaded_checkpoint_hash"] = checkpoint_hash
                        load_project(data)
                        # load_project may repair unsafe checkpoints and adjust the stage.
                        # Do not force the raw saved stage back after repair.
                        st.session_state.last_checkpoint_summary = loaded_project_summary()
                    st.success(st.session_state.last_checkpoint_summary)
                    st.rerun()
                else:
                    # Do not reload the same uploaded JSON on every Streamlit rerun.
                    # Otherwise every button click is immediately overwritten by the old checkpoint,
                    # making the UI look like buttons do not respond.
                    st.caption("Checkpoint already loaded. Continue using the app normally.")
            except Exception as exc:
                st.error(f"Could not load checkpoint: {exc}")

render_mobile_top_nav()
render_mobile_dashboard()

# =============================================================================
# Stage 1: Project Setup
# =============================================================================
if st.session_state.stage == 1:
    st.markdown('<div class="stage-header">🧭 Project Setup / 项目设置</div>', unsafe_allow_html=True)
    st.markdown('<div class="stage-desc">Create or load a reusable novel project. This app saves project data as JSON so you can continue later.</div>', unsafe_allow_html=True)

    tab_new, tab_load, tab_cloud, tab_backup = st.tabs(["Create / Edit Project", "Load Project JSON", "Cloud Save / Load", "Backup"])

    with tab_new:
        col1, col2 = st.columns(2)
        with col1:
            normalize_widget_defaults()
            st.text_input("Project file name", key="project_name", on_change=mark_dirty)
            st.text_input("Novel title / 书名", key="novel_title", on_change=mark_dirty)
            languages = ["Chinese 中文", "English", "Bilingual 中英双语"]
            st.selectbox("Writing language", languages, key="language", on_change=mark_dirty)
            audience_options = ["女频", "男频", "通用/无固定频道"]
            st.radio(
                "频道方向 / Audience",
                audience_options,
                key="audience_channel",
                horizontal=True,
                on_change=mark_dirty,
            )
            platform_options = ["起点中文网", "晋江文学城", "番茄小说", "七猫小说", "掌阅", "红袖添香", "潇湘书院", "知乎盐选 / 短篇付费", "飞卢", "刺猬猫 / 轻小说", "自定义/多平台"]
            selected_list_is_valid("publishing_platforms", platform_options, ["晋江文学城", "番茄小说"])
            st.multiselect(
                "目标投稿平台 / Publishing target",
                platform_options,
                key="publishing_platforms",
                help="建议主选 1 个平台，也可以多选用于比较不同读者口味。",
                on_change=mark_dirty,
            )
            styles = ["现代网文", "古风/仙侠", "女频仙侠", "玄幻升级流", "悬疑推理", "都市情感", "轻小说", "宫斗权谋", "中式恐怖", "文学感", "自定义"]
            st.selectbox(
                "Chinese style preset", styles, key="chinese_style", on_change=mark_dirty
            )
        with col2:
            genres = ["Xianxia / 仙侠", "Wuxia / 武侠", "Fantasy", "Science Fiction", "Romance", "Thriller", "Mystery", "Historical Fiction", "Urban Fantasy / 都市异能", "Horror", "Literary Fiction", "Other"]
            normalize_widget_defaults()
            st.selectbox("Primary genre / 主类型", genres, key="genre", on_change=mark_dirty)
            tag_options = ["重生", "复仇", "末世", "系统", "大女主", "强强", "逆袭", "爽文", "甜宠", "虐恋", "权谋", "宫斗", "悬疑", "刑侦", "法律诉讼", "异能", "修仙", "灵气复苏", "无限流", "穿书", "年代文", "克苏鲁", "种田", "经营", "校园", "娱乐圈"]
            selected_list_is_valid("genre_tags", tag_options, ["重生", "复仇", "大女主"])
            st.multiselect(
                "可叠加题材标签 / Multi-genre tags",
                tag_options,
                key="genre_tags",
                help="例如：重生 + 复仇 + 末世 + 系统 + 大女主",
                on_change=mark_dirty,
            )
            st.text_input("Custom tags / 自定义题材", key="custom_genre_tags", placeholder="例如：赛博修仙、律政复仇、末世经营", on_change=mark_dirty)
            submission_options = ["长篇连载", "短篇付费", "免费爽文", "签约向商业文", "IP 改编潜力文", "剧情流", "感情流", "升级流", "复仇流", "群像权谋", "单元案件流"]
            selected_list_is_valid("submission_types", submission_options, ["长篇连载", "签约向商业文"])
            st.multiselect(
                "投稿类型 / Article type",
                submission_options,
                key="submission_types",
                on_change=mark_dirty,
            )
            goal_options = ["前三章留存", "签约通过率", "追读率", "付费转化", "爽点密度", "人物关系张力", "IP潜力", "降低写崩风险"]
            selected_list_is_valid("platform_goals", goal_options, ["前三章留存", "签约通过率", "追读率"])
            st.multiselect(
                "平台目标 / Strategy goals",
                goal_options,
                key="platform_goals",
                on_change=mark_dirty,
            )
            lengths = ["Short story (~5 chapters)", "Novella (~10 chapters)", "Novel (~18 chapters)", "Long web novel (~30 chapters)"]
            st.selectbox(
                "Target length", lengths, key="target_length", on_change=mark_dirty
            )
            st.text_input("Style / tone note", key="style_note", placeholder="例如：古风但不堆辞藻，节奏快，对白自然", on_change=mark_dirty)

        st.text_area(
            "Story idea / 故事想法",
            key="story_idea",
            height=ui_height(130, 110),
            placeholder="例如：一个被逐出宗门的少女，在废弃灵矿中发现会说话的古剑……",
            on_change=mark_dirty,
        )
        st.text_area(
            "Global writing instruction / 全局写作要求",
            key="global_instruction",
            height=ui_height(90, 80),
            placeholder="例如：用中文写作；女主冷静克制；少解释，多用动作和对白推进剧情；避免翻译腔。",
            on_change=mark_dirty,
        )
        st.text_area(
            "自定义投稿/平台要求",
            key="custom_platform_note",
            height=80,
            placeholder="例如：主投番茄，第一章必须重生醒来+立刻危机+小反杀；前三章不要大段设定。",
            on_change=mark_dirty,
        )
        strategy_preview = platform_strategy_text()
        if strategy_preview:
            with st.expander("查看自动平台策略 / Platform strategy preview", expanded=False):
                st.markdown(strategy_preview.replace("\n", "  \n"))
        st.text_area(
            "Banned / overused phrases to avoid / 避免使用的套话",
            key="banned_phrases",
            height=ui_height(110, 90),
            on_change=mark_dirty,
        )
        st.info("Recommended workflow: Project Setup → Story Bible → Outline → Chapter Planner → Draft Writer → Revision Studio → Consistency Editor → Polish → Export.")
        if st.button("Save settings and go to Story Bible →", type="primary"):
            st.session_state.stage = 2
            st.rerun()

    with tab_load:
        if st.session_state.get("last_checkpoint_summary"):
            st.success(st.session_state.last_checkpoint_summary)
        uploaded = st.file_uploader("Upload project JSON", type=["json"])
        if uploaded is not None:
            try:
                with st.spinner("正在读取项目 JSON / Loading project JSON..."):
                    data = json.loads(uploaded.read().decode("utf-8"))
                    load_project(data)
                    st.session_state.last_checkpoint_summary = loaded_project_summary()
                st.success(st.session_state.last_checkpoint_summary)
                st.rerun()
            except Exception as exc:
                st.error(f"Could not load project: {exc}")

    with tab_cloud:
        st.markdown("### ☁️ Cloud Save / Load")
        render_cloud_projects_panel("setup_cloud")

    with tab_backup:
        backup = json.dumps(project_data(), ensure_ascii=False, indent=2)
        st.download_button(
            "Download project backup JSON",
            data=backup,
            file_name=f"{safe_filename(st.session_state.project_name)}.json",
            mime="application/json",
            use_container_width=True,
        )

# =============================================================================
# Stage 2: Story Bible
# =============================================================================
elif st.session_state.stage == 2:
    st.markdown('<div class="stage-header">📚 Story Bible / 小说设定集</div>', unsafe_allow_html=True)
    st.markdown('<div class="stage-desc">Keep names, relationships, rules, timeline, locations, clues, and style constraints consistent across chapters.</div>', unsafe_allow_html=True)

    col1, col2 = st.columns(2)
    with col1:
        st.text_area("Characters / 人物设定", key="characters", height=220, on_change=mark_dirty)
        st.text_area("Timeline / 时间线", key="timeline", height=180, on_change=mark_dirty)
    with col2:
        st.text_area("Worldbuilding, rules, factions / 世界观、规则、势力", key="worldbuilding", height=220, on_change=mark_dirty)
        st.text_area("Glossary: names, terms, objects / 术语表", key="glossary", height=180, on_change=mark_dirty)

    st.text_area(
        "Full Story Bible / 总设定集",
        key="story_bible",
        height=260,
        placeholder="核心设定、人物关系、伏笔、禁忌、世界规则、写作风格规则……",
        on_change=mark_dirty,
    )

    col_a, col_b = st.columns(2)
    with col_a:
        if st.button("Generate / improve Story Bible ✨", type="primary", use_container_width=True):
            with st.spinner("Building Story Bible..."):
                placeholder = st.empty()
                result = generate_text(
                    system="你是资深小说设定架构师。请直接产出可用于长篇连载的中文《故事设定集》。不要输出分析过程、提示词、假设、检查清单或追问。",
                    user=f"""Writing context:
{writing_context()}

Story idea:
{st.session_state.story_idea}

Existing character notes:
{st.session_state.characters}

Existing worldbuilding notes:
{st.session_state.worldbuilding}

Existing timeline notes:
{st.session_state.timeline}

Existing glossary:
{st.session_state.glossary}

请直接生成成品《故事设定集》。第一行必须是：
《{st.session_state.novel_title}》故事设定集

禁止输出：提示词、分析过程、英文解释、检查清单、自我修正、写作计划、追问。

必须使用以下中文章节标题：
## 核心前提
## 主要角色
## 人物关系
## 世界法则与力量体系
## 核心地点与势力
## 时间线
## 重要道具 / 线索 / 伏笔
## 文风准则
## 必须避免
""",
                    placeholder=placeholder,
                    max_tokens=8000,
                    task_name="Story Bible generation",
                )
            if result:
                st.session_state["_last_story_bible_raw"] = st.session_state.get("_last_ai_raw_output", result)
                result = clean_story_bible_output(result, st.session_state.novel_title)
                placeholder.markdown(result)
                st.session_state.story_bible = result
                mark_dirty()
                if story_bible_is_incomplete(result):
                    st.warning("Story Bible may be incomplete. Use Continue / 补全 below before moving on.")
                cloud_auto_save("Story Bible generation")
                st.success("Story Bible updated.")
    with col_b:
        custom = st.text_area("修改要求（只给 AI 看，不会写入设定集）", height=90, key="story_bible_revision_instruction", placeholder="例如：增加修炼等级体系；把男主身份隐藏更久；加入三条伏笔。")
        if st.button("Revise Story Bible 🔁", use_container_width=True):
            if not st.session_state.story_bible:
                st.error("Generate or enter a Story Bible first.")
            elif not custom.strip():
                st.error("Enter a revision instruction.")
            else:
                with st.spinner("Revising Story Bible..."):
                    placeholder = st.empty()
                    result = generate_text(
                        system="你是资深小说设定架构师。请直接返回修订后的中文《故事设定集》。不要输出分析过程、提示词、假设、检查清单或追问。",
                        user=f"""Writing context:
{writing_context()}

Current Story Bible:
{st.session_state.story_bible}

Revision instruction:
{custom}

请直接返回修订后的成品设定集。不要显示提示词、分析过程、修改说明或追问。
""",
                        placeholder=placeholder,
                        max_tokens=8000,
                        task_name="Story Bible revision",
                    )
                if result:
                    st.session_state["_last_story_bible_raw"] = st.session_state.get("_last_ai_raw_output", result)
                    result = clean_story_bible_output(result, st.session_state.novel_title)
                    placeholder.markdown(result)
                    st.session_state.story_bible = result
                    mark_dirty()
                    if story_bible_is_incomplete(result):
                        st.warning("Story Bible may be incomplete. Use Continue / 补全 below before moving on.")
                    cloud_auto_save("Story Bible revision")
                    st.success("Story Bible revised.")

    with st.expander("🧯 Story Bible recovery / 补全与恢复", expanded=story_bible_is_incomplete(st.session_state.get("story_bible", ""))):
        if story_bible_is_incomplete(st.session_state.get("story_bible", "")):
            st.warning("当前 Story Bible 可能没有生成完整，建议点击补全。")
        if st.session_state.get("_last_story_bible_raw"):
            if st.button("Restore last raw AI output / 恢复最近原始输出", use_container_width=True):
                st.session_state.story_bible = st.session_state.get("_last_story_bible_raw", "")
                mark_dirty()
                st.rerun()
        if st.button("Continue / 补全 Story Bible ✨", type="primary", use_container_width=True):
            with st.spinner("Continuing Story Bible..."):
                placeholder = st.empty()
                result = generate_text(
                    system="你是资深小说设定架构师。请只补全当前设定集中缺失或过短的部分，不要重写已经完整的部分。",
                    user=f"""Writing context:
{writing_context()}

Current Story Bible, possibly incomplete:
{st.session_state.story_bible}

请补全缺失部分，尤其检查并补全这些章节：
## 核心前提
## 主要角色
## 人物关系
## 世界法则与力量体系
## 核心地点与势力
## 时间线
## 重要道具 / 线索 / 伏笔
## 文风准则
## 必须避免

请返回完整的修订后《故事设定集》，不要输出说明。
""",
                    placeholder=placeholder,
                    max_tokens=8000,
                    task_name="Story Bible continuation",
                )
            if result:
                st.session_state["_last_story_bible_raw"] = st.session_state.get("_last_ai_raw_output", result)
                result = clean_story_bible_output(result, st.session_state.novel_title)
                st.session_state.story_bible = result
                mark_dirty()
                cloud_auto_save("Story Bible continuation")
                st.success("Story Bible continued/completed.")
                st.rerun()

    if st.button("Next: Outline →", type="primary"):
        st.session_state.stage = 3
        st.rerun()

# =============================================================================
# Stage 3: Outline
# =============================================================================
elif st.session_state.stage == 3:
    st.markdown('<div class="stage-header">📝 Outline / 大纲</div>', unsafe_allow_html=True)
    st.markdown('<div class="stage-desc">Generate or repair a chapter-by-chapter outline grounded in the Story Bible.</div>', unsafe_allow_html=True)

    ensure_chapters_available()

    if st.session_state.get("_json_repair_note"):
        st.info(st.session_state.get("_json_repair_note"))

    current_chapter_count = len(st.session_state.get("chapters", []) or [])
    has_outline = bool((st.session_state.get("outline") or "").strip())
    has_story_bible = bool((st.session_state.get("story_bible") or "").strip())

    status_cols = st.columns(4)
    status_cols[0].metric("Story Bible", "Yes" if has_story_bible else "No")
    status_cols[1].metric("Outline", "Yes" if has_outline else "No")
    status_cols[2].metric("Chapters", current_chapter_count)
    status_cols[3].metric("Target", desired_chapter_count())

    if has_story_bible and not has_outline:
        st.warning("这个项目有 Story Bible，但还没有正式大纲。你可以用 AI 生成正式大纲，也可以先用 Story Bible 时间线继续。")
    elif has_outline and current_chapter_count == 0:
        st.warning("当前有大纲文本，但还没有识别到章节。请重新识别、手动改成 第1章：标题 — 摘要，或重新生成正式大纲。")

    action_cols = st.columns([0.35, 0.22, 0.22, 0.21])
    with action_cols[0]:
        gen_clicked = st.button("Generate formal outline ✨", type="primary", use_container_width=True, key="outline_generate_v12")
    with action_cols[1]:
        local_clicked = st.button("Build outline locally", use_container_width=True, key="outline_local_v12")
    with action_cols[2]:
        detect_clicked = st.button("Re-detect chapters", use_container_width=True, key="outline_detect_v12")
    with action_cols[3]:
        continue_clicked = st.button("Continue →", use_container_width=True, key="outline_continue_v12")

    if local_clicked:
        set_last_action("Building local outline from available Story Bible/chapters. No AI call used.")
        st.session_state.outline = build_local_outline_from_material()
        st.session_state.chapters = parse_chapters(st.session_state.outline)
        st.session_state["_json_repair_note"] = ""
        cloud_auto_save("local outline build")
        st.success(f"已本地生成可用大纲，识别到 {len(st.session_state.chapters)} 个章节。")
        st.rerun()

    if detect_clicked:
        set_last_action("Re-detect chapters clicked.")
        count_before = len(st.session_state.get("chapters", []) or [])
        count_after = ensure_chapters_available()
        if count_after:
            st.success(f"已识别 {count_after} 个章节。")
        else:
            st.warning("仍然没有识别到章节。可以点击 Build outline locally 先生成可用大纲，或点击 Generate formal outline 调用 AI。")
        set_last_action(f"Re-detect complete: {count_before} -> {count_after} chapters.")

    if continue_clicked:
        set_last_action("Continue from Outline clicked.")
        if ensure_chapters_available():
            st.session_state.stage = 4
            st.rerun()
        else:
            st.error("当前没有可用章节，不能进入 Chapter Planner。请先点击 Build outline locally 或 Generate formal outline。")

    if gen_clicked:
        set_last_action("Generate formal outline clicked. Preparing AI request...")
        source_text = (st.session_state.get("story_bible") or "").strip() or (st.session_state.get("story_idea") or "").strip()
        if not source_text:
            st.error("没有 Story Bible 或 Story Idea。请先在 Project Setup / Story Bible 填写内容。")
            set_last_action("Generate outline stopped: no Story Bible or Story Idea.")
        else:
            st.info("已触发 Generate formal outline，正在调用 AI。若失败，错误会显示在这里和左侧 Last action。")
            with st.spinner("Generating formal outline with AI..."):
                placeholder = st.empty()
                result = generate_text(
                    system="You are a professional novel editor and story architect. Create a detailed, usable chapter outline. Return only final content, no notes.",
                    user=f"""Writing context:
{writing_context()}

Create a clean formal chapter-by-chapter outline for about {desired_chapter_count()} chapters.
Use the Story Bible as the authority. If the current outline is temporary or recovered from a timeline, rewrite it into a formal outline.

Source material:
{source_text}

Current outline, if any:
{st.session_state.get('outline') or 'None'}

Required format exactly:
# {st.session_state.get('novel_title', 'Untitled')}

## Characters
[2-6 main characters with brief descriptions]

## Story Arc
[Main conflict, escalation, midpoint, climax, ending direction]

## Chapter Outline
第1章：标题 — 2-3句中文摘要。
第2章：标题 — 2-3句中文摘要。
...continue for all chapters

## Foreshadowing Map
[List major clues and when they should pay off]
""",
                    placeholder=placeholder,
                    max_tokens=5000,
                )
            if result:
                st.session_state.outline = result
                st.session_state.chapters = parse_chapters(result)
                mark_dirty()
                cloud_auto_save("outline generation")
                st.session_state["_json_repair_note"] = ""
                if st.session_state.chapters:
                    set_last_action(f"Formal outline generated: {len(st.session_state.chapters)} chapters detected.")
                    st.success(f"正式大纲已生成，识别到 {len(st.session_state.chapters)} 个章节。")
                    st.rerun()
                else:
                    set_last_action("AI returned outline but parser detected 0 chapters.")
                    st.warning("AI 返回了大纲，但没有识别到章节。你可以点击 Build outline locally 或手动把章节改成：第1章：标题 — 摘要。")
            else:
                set_last_action("AI returned empty result for outline. Use Build outline locally to continue.")
                st.warning("AI 没有返回结果。可以先点击 Build outline locally 继续项目，或检查左侧 AI connection check。")

    st.divider()
    if st.session_state.get("outline"):
        st.subheader(st.session_state.get("novel_title", "Untitled"))
        view_tab, raw_tab = st.tabs(["章节分段视图", "原始大纲"] )
        with view_tab:
            render_formatted_outline()
        with raw_tab:
            st.markdown(st.session_state.outline)
    else:
        st.info("当前还没有大纲。点击 Generate formal outline 或 Build outline locally。")

    st.divider()
    revision = st.text_area("Revise outline instruction", height=90, key="outline_revision_instruction", placeholder="例如：增加反派动机；减少爱情线；每三章有一个小高潮。")
    rev_cols = st.columns([0.35, 0.65])
    with rev_cols[0]:
        revise_clicked = st.button("Revise outline 🔁", use_container_width=True, key="outline_revise_v12")
    if revise_clicked:
        if not (st.session_state.get("outline") or "").strip():
            st.error("没有可修订的大纲。请先生成大纲。")
        elif not revision.strip():
            st.error("Enter a revision instruction.")
        else:
            set_last_action("Revise outline clicked. Calling AI...")
            with st.spinner("Revising outline..."):
                placeholder = st.empty()
                result = generate_text(
                    system="You are a story architect. Revise the outline while preserving clear chapter format. Return only the revised outline.",
                    user=f"""Writing context:
{writing_context()}

Current outline:
{st.session_state.outline}

Revision instruction:
{revision}

Return the revised outline in this format:
第1章：标题 — 摘要
第2章：标题 — 摘要
""",
                    placeholder=placeholder,
                    max_tokens=5000,
                )
            if result:
                st.session_state.outline = result
                st.session_state.chapters = parse_chapters(result)
                mark_dirty()
                cloud_auto_save("outline revision")
                set_last_action(f"Outline revised: {len(st.session_state.chapters)} chapters detected.")
                st.success(f"Outline revised. {len(st.session_state.chapters)} chapters detected.")
                st.rerun()

    if st.button("Next: Chapter Planner →", type="primary", key="outline_next_v12"):
        if ensure_chapters_available():
            st.session_state.stage = 4
            st.rerun()
        else:
            st.error("没有章节，不能进入 Chapter Planner。请先生成或本地构建大纲。")

# =============================================================================
# Stage 4: Chapter Planner
# =============================================================================
elif st.session_state.stage == 4:
    st.markdown('<div class="stage-header">🧩 Chapter Planner / 章节规划</div>', unsafe_allow_html=True)
    st.markdown('<div class="stage-desc">Create beat sheets before drafting. You can plan one chapter or generate missing beat sheets for the whole book.</div>', unsafe_allow_html=True)

    if not st.session_state.chapters:
        st.warning("Generate an outline first.")
        st.stop()

    total_chapters = len(st.session_state.chapters)
    planned_count = len([i for i in range(total_chapters) if st.session_state.chapter_beats.get(i)])
    st.info(f"Detected {total_chapters} chapters from the outline. Beat sheets ready: {planned_count}/{total_chapters}.")

    if total_chapters <= 1:
        st.warning(
            "Only one chapter was detected. If your outline should have more chapters, go back to Outline and revise/regenerate it. "
            "The parser is now more tolerant, but the outline still needs recognizable chapter headings like 第1章/Chapter 1/1."
        )

    with st.expander("📚 Whole-book beat sheet generation / 全书章节规划", expanded=True):
        st.caption("Safest default: generate only missing beat sheets, so your hand-edited chapter plans are not overwritten.")
        whole_cols = st.columns([0.35, 0.35, 0.30])
        with whole_cols[0]:
            overwrite_beats = st.checkbox("Overwrite existing beat sheets", value=False, help="默认不覆盖已有章节规划。勾选后会重写已有 Beat Sheet。")
        with whole_cols[1]:
            max_run = st.selectbox("Max chapters this run", ["5", "10", "All"], index=0)
        with whole_cols[2]:
            plan_style = st.selectbox("Planning depth", ["Practical", "Detailed", "Very detailed"], index=0)

        missing_indices = [i for i in range(total_chapters) if overwrite_beats or not st.session_state.chapter_beats.get(i)]
        if max_run != "All":
            run_indices = missing_indices[: int(max_run)]
        else:
            run_indices = missing_indices

        st.caption(f"This run will plan {len(run_indices)} chapter(s). Remaining/missing candidates: {len(missing_indices)}.")

        if st.button("Generate missing beat sheets for whole book ✨", type="primary", use_container_width=True):
            if not run_indices:
                st.success("All chapters already have beat sheets. Check 'Overwrite existing beat sheets' if you want to regenerate them.")
            else:
                progress = st.progress(0)
                status = st.empty()
                for pos, i in enumerate(run_indices, 1):
                    ch_i = st.session_state.chapters[i]
                    status.info(f"Planning {pos}/{len(run_indices)}: {chapter_heading(ch_i)}")
                    placeholder = st.empty()
                    result = generate_text(
                        system=(
                            "You are a chapter-level story planner for commercial web fiction. "
                            "Create practical beat sheets that preserve continuity, support strong pacing, and end with a hook. "
                            "Return only the beat sheet content."
                        ),
                        user=f"""Writing context:
{writing_context()}

Overall outline:
{st.session_state.outline[:7000]}

Chapter to plan:
{chapter_heading(ch_i)}
Summary: {ch_i.get('summary', '')}

Planning depth: {plan_style}

Create a chapter beat sheet with:
## 本章定位
## POV / 视角人物
## 场景地点
## 本章目标
## 外部冲突
## 内在/情绪冲突
## 本章爽点或情绪价值
## 新信息/伏笔
## 人物变化
## 必须包含
## 必须避免
## 断章钩子
## 写作提醒
""",
                        placeholder=placeholder,
                        max_tokens=2500,
                    )
                    if result:
                        st.session_state.chapter_beats[i] = result
                        cloud_auto_save(f"beat sheet {chapter_heading(ch_i)}")
                    progress.progress(pos / len(run_indices))
                status.success(f"Generated {len(run_indices)} beat sheet(s).")
                st.rerun()

    st.divider()
    st.subheader("Single chapter planner / 单章规划")
    labels = [f"{chapter_heading(ch)} {'✓' if i in st.session_state.chapter_beats else ''}" for i, ch in enumerate(st.session_state.chapters)]
    idx = st.selectbox("Select chapter", range(len(labels)), format_func=lambda i: labels[i])
    ch = st.session_state.chapters[idx]
    st.markdown(f"**Outline summary:** {ch.get('summary', '')}")

    current = st.session_state.chapter_beats.get(idx, "")
    beat_key = f"chapter_beat_editor_{idx}"
    beat_hash_key = f"chapter_beat_editor_hash_{idx}"
    current_hash = hashlib.md5(current.encode("utf-8")).hexdigest() if current else ""
    if beat_key not in st.session_state or st.session_state.get(beat_hash_key) != current_hash:
        st.session_state[beat_key] = current
        st.session_state[beat_hash_key] = current_hash
    edited_beat = st.text_area("Chapter beat sheet", key=beat_key, height=ui_height(260, 220), placeholder="POV, location, goal, conflict, reveal, emotional turn, ending hook, must include, must avoid...")
    if edited_beat != current:
        st.session_state.chapter_beats[idx] = edited_beat
        st.session_state[beat_hash_key] = hashlib.md5(edited_beat.encode("utf-8")).hexdigest() if edited_beat else ""
        mark_dirty()

    col1, col2 = st.columns(2)
    with col1:
        if st.button("Generate selected chapter beat sheet ✨", type="primary", use_container_width=True):
            with st.spinner("Planning chapter..."):
                placeholder = st.empty()
                result = generate_text(
                    system="You are a chapter-level story planner. Create practical beat sheets that help draft strong scenes. Return only the beat sheet.",
                    user=f"""Writing context:
{writing_context()}

Overall outline:
{st.session_state.outline[:5000]}

Chapter to plan:
{chapter_heading(ch)}
Summary: {ch.get('summary', '')}

Create a chapter beat sheet with:
## 本章定位
## POV / 视角人物
## 场景地点
## 本章目标
## 外部冲突
## 内在/情绪冲突
## 本章爽点或情绪价值
## 新信息/伏笔
## 人物变化
## 必须包含
## 必须避免
## 断章钩子
## 写作提醒
""",
                    placeholder=placeholder,
                    max_tokens=2500,
                )
            if result:
                st.session_state.chapter_beats[idx] = result
                cloud_auto_save(f"beat sheet {chapter_heading(ch)}")
                st.success("Beat sheet generated.")
                st.rerun()
    with col2:
        instruction = st.text_area("Beat revision instruction", height=90, key=f"beat_revision_instruction_{idx}", placeholder="例如：结尾更强；增加女主隐藏实力；减少男主出场。")
        if st.button("Revise selected beat sheet 🔁", use_container_width=True):
            if not st.session_state.chapter_beats.get(idx):
                st.error("Generate or enter a beat sheet first.")
            elif not instruction.strip():
                st.error("Enter a revision instruction.")
            else:
                with st.spinner("Revising beat sheet..."):
                    placeholder = st.empty()
                    result = generate_text(
                        system="You are a chapter planner. Revise the beat sheet clearly and practically. Return only the revised beat sheet.",
                        user=f"""Writing context:
{writing_context()}

Current beat sheet:
{st.session_state.chapter_beats[idx]}

Revision instruction:
{instruction}

Return the revised beat sheet.
""",
                        placeholder=placeholder,
                        max_tokens=2500,
                    )
                if result:
                    st.session_state.chapter_beats[idx] = result
                    cloud_auto_save(f"beat sheet revision {chapter_heading(ch)}")
                    st.success("Beat sheet revised.")
                    st.rerun()

    if st.button("Next: Draft Writer →", type="primary"):
        st.session_state.stage = 5
        st.rerun()

# =============================================================================
# Stage 5: Draft Writer
# =============================================================================
elif st.session_state.stage == 5:
    st.markdown('<div class="stage-header">✍️ Draft Writer / 章节写作</div>', unsafe_allow_html=True)
    if is_mobile_mode():
        st.info("Mobile mode: 当前页优先用于继续写、改当前章、快速保存。复杂设定建议回电脑端处理。")
    st.markdown('<div class="stage-desc">Draft chapters from outline + Story Bible + beat sheet. Revisions are proposed side-by-side before acceptance.</div>', unsafe_allow_html=True)

    if not st.session_state.chapters:
        st.warning("Generate an outline first.")
        st.stop()

    labels = [f"{chapter_heading(ch)} {'✓' if i in st.session_state.chapter_drafts else ''}" for i, ch in enumerate(st.session_state.chapters)]
    idx = st.selectbox("Select chapter", range(len(labels)), format_func=lambda i: labels[i])
    ch = st.session_state.chapters[idx]
    beat = st.session_state.chapter_beats.get(idx, "")

    st.markdown(f"**Outline summary:** {ch.get('summary', '')}")
    render_quick_reference(idx)
    render_generate_full_text_panel("draft_writer")
    if beat:
        with st.expander("View beat sheet"):
            st.markdown(beat)
    else:
        st.warning("This chapter has no beat sheet yet. You can still draft, but quality may be lower.")

    length_target = st.selectbox("Length target", ["Short scene", "Normal chapter", "Long chapter"], index=1)
    length_rule = {
        "Short scene": "700-1000 English words or 1000-1800 Chinese characters",
        "Normal chapter": "1000-1600 English words or 1800-3000 Chinese characters",
        "Long chapter": "1600-2400 English words or 3000-4500 Chinese characters",
    }[length_target]

    if st.button(f"Draft {chapter_heading(ch)} ✨", type="primary"):
        with st.spinner("Drafting chapter..."):
            placeholder = st.empty()
            result = generate_chapter_text(idx, length_rule, placeholder)
        if result:
            if st.session_state.chapter_drafts.get(idx):
                push_version(idx, st.session_state.chapter_drafts[idx], "Before new draft")
            st.session_state.chapter_drafts[idx] = result
            rebuild_full_draft()
            mark_dirty()
            cloud_auto_save(f"draft {chapter_heading(ch)}")
            st.success("Chapter drafted.")
            st.rerun()

    if st.session_state.chapter_drafts.get(idx):
        st.subheader("Current draft")
        draft_key = f"draft_editor_{idx}"
        draft_hash_key = f"draft_editor_hash_{idx}"
        current_draft = st.session_state.chapter_drafts[idx]
        current_draft_hash = hashlib.md5(current_draft.encode("utf-8")).hexdigest() if current_draft else ""
        if draft_key not in st.session_state or st.session_state.get(draft_hash_key) != current_draft_hash:
            st.session_state[draft_key] = current_draft
            st.session_state[draft_hash_key] = current_draft_hash
        edited_draft = st.text_area("Draft text", key=draft_key, height=ui_height(360, 280))
        if edited_draft != current_draft:
            st.info("Draft has local edits. They are saved into this project state automatically; use cloud save to persist them online.")
            st.session_state.chapter_drafts[idx] = edited_draft
            st.session_state[draft_hash_key] = hashlib.md5(edited_draft.encode("utf-8")).hexdigest() if edited_draft else ""
            rebuild_full_draft()
            mark_dirty()
        render_batch_revision_panel("draft_writer_current")
        render_local_refinement_panel(idx, st.session_state.chapter_drafts[idx], f"draft_{idx}")
        with st.expander("📊 给当前章节打分 / Quick market score", expanded=False):
            render_market_review_controls(f"draft_quick_{idx}", compact=True, chapter_idx=idx)

        st.subheader("Revise this chapter")
        mode = st.selectbox("Quick revision mode", [
            "Make it more suspenseful", "Make dialogue more natural", "Add Chinese web-novel pacing",
            "Make prose less AI-like", "Increase emotional tension", "Add sensory detail", "Reduce exposition",
            "Add conflict", "Make ending stronger", "Rewrite in 古风 tone",
        ])
        custom = st.text_area("Additional revision instruction", height=90, key=f"draft_revision_instruction_{idx}", placeholder="例如：让女主少说话，用动作表现她的防备。")
        if st.button("Create revision proposal 🔁"):
            instruction = revision_prompt_for_mode(mode, custom)
            with st.spinner("Creating revision proposal..."):
                placeholder = st.empty()
                result = generate_text(
                    system="You are a professional fiction editor. Create a revised version without losing useful details.",
                    user=f"""Writing context:
{writing_context()}

Chapter:
{chapter_heading(ch)}

Current draft:
{st.session_state.chapter_drafts[idx]}

Revision instruction:
{instruction}

Return the revised chapter text only.
""",
                    placeholder=placeholder,
                    max_tokens=5000,
                )
            if result:
                st.session_state.pending_revisions[idx] = result
                cloud_auto_save(f"revision studio proposal {chapter_heading(ch)}")
                st.success("Revision proposal created.")
                st.rerun()

        if st.session_state.pending_revisions.get(idx):
            st.subheader("Compare before accepting")
            left, right = st.columns(2)
            with left:
                st.markdown("**Original**")
                st.text_area("Original", value=st.session_state.chapter_drafts[idx], height=360, key=f"orig_{idx}")
            with right:
                st.markdown("**Proposed revision**")
                st.text_area("Proposed", value=st.session_state.pending_revisions[idx], height=360, key=f"prop_{idx}")
            a, b, c = st.columns(3)
            with a:
                if st.button("Accept revision", type="primary"):
                    push_version(idx, st.session_state.chapter_drafts[idx], "Before accepted revision")
                    st.session_state.chapter_drafts[idx] = st.session_state.pending_revisions.pop(idx)
                    rebuild_full_draft()
                    cloud_auto_save(f"accepted revision {chapter_heading(ch)}")
                    st.success("Revision accepted.")
                    st.rerun()
            with b:
                if st.button("Save proposal as version"):
                    push_version(idx, st.session_state.pending_revisions[idx], "Saved proposal")
                    cloud_auto_save(f"saved proposal version {chapter_heading(ch)}")
                    st.success("Proposal saved in version history.")
            with c:
                if st.button("Discard proposal"):
                    st.session_state.pending_revisions.pop(idx, None)
                    st.rerun()

    if st.button("Next: Revision Studio →", type="primary"):
        st.session_state.stage = 6
        st.rerun()

# =============================================================================
# Stage 6: Revision Studio
# =============================================================================
elif st.session_state.stage == 6:
    st.markdown('<div class="stage-header">🔁 Revision Studio / 修改工作台</div>', unsafe_allow_html=True)
    st.markdown('<div class="stage-desc">Continue writing, rewrite passages, score chapter quality, and manage version history.</div>', unsafe_allow_html=True)

    if not st.session_state.chapter_drafts:
        st.warning("Draft at least one chapter first.")
        st.stop()

    drafted_indices = sorted(st.session_state.chapter_drafts.keys())
    idx = st.selectbox("Select drafted chapter", drafted_indices, format_func=lambda i: chapter_heading(st.session_state.chapters[i]) if i < len(st.session_state.chapters) else f"Chapter {i+1}")
    ch = st.session_state.chapters[idx] if idx < len(st.session_state.chapters) else {"num": idx + 1, "title": f"Chapter {idx + 1}", "summary": ""}
    draft = st.session_state.chapter_drafts[idx]
    render_quick_reference(idx)
    render_generate_full_text_panel("revision_studio")
    render_batch_revision_panel("revision_studio")
    rev_draft_key = f"revision_current_editor_{idx}"
    rev_hash_key = f"revision_current_hash_{idx}"
    draft_hash = hashlib.md5(draft.encode("utf-8")).hexdigest() if draft else ""
    if rev_draft_key not in st.session_state or st.session_state.get(rev_hash_key) != draft_hash:
        st.session_state[rev_draft_key] = draft
        st.session_state[rev_hash_key] = draft_hash
    edited_revision_draft = st.text_area("Current chapter draft", key=rev_draft_key, height=ui_height(320, 260))
    if edited_revision_draft != draft:
        st.session_state.chapter_drafts[idx] = edited_revision_draft
        st.session_state[rev_hash_key] = hashlib.md5(edited_revision_draft.encode("utf-8")).hexdigest() if edited_revision_draft else ""
        rebuild_full_draft()
        mark_dirty()
        draft = edited_revision_draft

    tab_continue, tab_local, tab_quality, tab_versions = st.tabs(["Continue / Rewrite", "局部点杀精修", "Quality Score", "Version History"])
    with tab_continue:
        continue_from = st.text_area("Paste last paragraph or selected passage", height=120, placeholder="Paste the end of the scene or a passage to continue/rewrite...")
        action = st.selectbox("Action", ["Continue 800 Chinese characters / 500 English words", "Continue until scene ending", "Rewrite selected passage", "Rewrite with more dialogue", "Rewrite with more action", "Rewrite with more emotional conflict"])
        extra = st.text_area("Extra instruction", height=90)
        if st.button("Run continuation / rewrite ✨", type="primary"):
            source = continue_from.strip() or draft[-1200:]
            with st.spinner("Writing..."):
                placeholder = st.empty()
                result = generate_text(
                    system="You are a fiction writer and line editor. Continue or rewrite prose seamlessly.",
                    user=f"""Writing context:
{writing_context()}

Current chapter context:
{draft[-2500:]}

Selected passage or continuation point:
{source}

Action: {action}
Extra instruction: {extra or 'None'}

Return only the new or rewritten prose.
""",
                    placeholder=placeholder,
                    max_tokens=3500,
                )
            if result:
                st.session_state.pending_revisions[idx] = result
                cloud_auto_save(f"revision studio proposal {chapter_heading(ch)}")
                st.success("Output created. Copy it into the chapter or accept it as a replacement in Draft Writer.")
                st.markdown(result)

    with tab_local:
        render_local_refinement_panel(idx, draft, f"revision_{idx}")

    with tab_quality:
        if st.button("Score chapter quality 📊"):
            with st.spinner("Scoring chapter..."):
                placeholder = st.empty()
                result = generate_text(
                    system="You are a strict but helpful fiction editor. Score chapter craft and give actionable fixes.",
                    user=f"""Writing context:
{writing_context()}

Chapter draft:
{draft[:6000]}

Score with 1-10 ratings:
- Hook strength
- Conflict
- Character voice
- Dialogue naturalness
- Pacing
- Emotional tension
- Continuity risk
- AI-sounding risk

Then give the top 5 concrete fixes.
""",
                    placeholder=placeholder,
                    max_tokens=2500,
                )
            if result:
                st.session_state.quality_reports[idx] = result
                cloud_auto_save(f"quality score {chapter_heading(ch)}")
        if st.session_state.quality_reports.get(idx):
            st.markdown(st.session_state.quality_reports[idx])

    with tab_versions:
        versions = st.session_state.chapter_versions.get(idx, [])
        if not versions:
            st.info("No saved versions yet. Versions are created when you accept revisions or save proposals.")
        for v_i, version in enumerate(reversed(versions), 1):
            with st.expander(f"Version {len(versions)-v_i+1}: {version.get('label')} — {version.get('time')}"):
                st.text_area("Version text", value=version.get("text", ""), height=220, key=f"version_{idx}_{v_i}")

    if st.button("Next: Web Novel Score →", type="primary"):
        st.session_state.stage = 7
        st.rerun()

# =============================================================================
# Stage 7: Web Novel Score
# =============================================================================
elif st.session_state.stage == 7:
    st.markdown('<div class="stage-header">📊 Web Novel Score / 网文评审团</div>', unsafe_allow_html=True)
    st.markdown('<div class="stage-desc">模拟不同平台读者、主编和作者续写风险评委，判断作品追读、签约、平台匹配和修改优先级。</div>', unsafe_allow_html=True)

    render_quick_reference(None)
    render_market_review_controls("market_stage", compact=False, chapter_idx=None, show_inline_report=False)

    st.divider()
    st.subheader("Historical scoring reports / 历史评分")
    if not st.session_state.market_reviews:
        st.info("还没有评分报告。选择评分对象和评委后点击开始评分。")
    else:
        for k, review in reversed(list(st.session_state.market_reviews.items())):
            with st.expander(f"{review.get('time', '')} — {review.get('score_target', '')} — {review.get('source_label', '')}"):
                st.caption(f"评委：{', '.join(review.get('judges', []))} | 语气：{review.get('tone', '')} | 平台：{review.get('target_platform', '')}")
                st.markdown(review.get("report", ""))

    if st.button("Next: Consistency Editor →", type="primary"):
        st.session_state.stage = 8
        st.rerun()

# =============================================================================
# Stage 8: Consistency Editor
# =============================================================================
elif st.session_state.stage == 8:
    st.markdown('<div class="stage-header">🔍 Consistency Editor / 连续性与设定检查</div>', unsafe_allow_html=True)
    st.markdown('<div class="stage-desc">Check manuscript against Story Bible: character facts, timeline, rules, clues, tone, and Chinese name consistency.</div>', unsafe_allow_html=True)

    rebuild_full_draft()
    draft = st.session_state.full_draft or st.session_state.outline
    if not draft:
        st.warning("Draft chapters first.")
        st.stop()

    render_quick_reference(None)
    render_generate_full_text_panel("consistency_editor")

    focus = st.multiselect(
        "Check types",
        ["Character consistency", "Timeline consistency", "Worldbuilding/rule consistency", "Relationship consistency", "Foreshadowing/clue consistency", "Chinese name consistency", "Tone consistency", "Actionable rewrite plan"],
        default=["Character consistency", "Timeline consistency", "Worldbuilding/rule consistency", "Foreshadowing/clue consistency", "Actionable rewrite plan"],
    )
    extra = st.text_area("Extra check instruction", height=80, placeholder="例如：重点检查修仙等级是否前后矛盾，伏笔是否太早泄露。")

    if st.button("Run consistency check 🔍", type="primary"):
        with st.spinner("Checking consistency..."):
            placeholder = st.empty()
            result = generate_text(
                system="You are a professional continuity editor for long-form fiction.",
                user=f"""Writing context:
{writing_context()}

Check types:
{', '.join(focus)}

Extra instruction:
{extra or 'None'}

Story Bible:
{compact_story_bible(5000)}

Outline:
{st.session_state.outline[:4000]}

Manuscript excerpt / current draft:
{draft[:9000]}

Return:
## High-priority continuity problems
## Possible contradictions with Story Bible
## Timeline / relationship / rule issues
## Foreshadowing and unresolved clues
## Specific chapter-level fixes
## Safe rewrite plan
""",
                placeholder=placeholder,
                max_tokens=4500,
            )
        if result:
            st.session_state.logic_report = result
            cloud_auto_save("consistency check")
            st.success("Consistency check complete.")

    if st.session_state.logic_report:
        st.markdown(st.session_state.logic_report)
        revision = st.text_area("Revise consistency report instruction", height=80, placeholder="例如：改成表格；只保留最重要的10个问题；给出逐章修复清单。")
        if st.button("Revise report 🔁"):
            if revision.strip():
                with st.spinner("Revising report..."):
                    placeholder = st.empty()
                    result = generate_text(
                        system="You are a continuity editor. Revise the report to be more useful.",
                        user=f"Current report:\n{st.session_state.logic_report}\n\nRevision instruction:\n{revision}",
                        placeholder=placeholder,
                        max_tokens=3500,
                    )
                if result:
                    st.session_state.logic_report = result
                    cloud_auto_save("consistency report revision")
                    st.rerun()

    if st.button("Next: Polish →", type="primary"):
        st.session_state.stage = 9
        st.rerun()

# =============================================================================
# Stage 9: Polish
# =============================================================================
elif st.session_state.stage == 9:
    st.markdown('<div class="stage-header">✨ Polish / 润色去AI味</div>', unsafe_allow_html=True)
    st.markdown('<div class="stage-desc">Polish selected chapters for natural prose, Chinese rhythm, dialogue, pacing, and style.</div>', unsafe_allow_html=True)

    rebuild_full_draft()
    if not st.session_state.full_draft:
        st.warning("Draft chapters first.")
        st.stop()

    render_quick_reference(None)
    render_generate_full_text_panel("polish")

    choices = ["Full draft excerpt"] + [chapter_heading(ch) for i, ch in enumerate(st.session_state.chapters) if i in st.session_state.chapter_drafts]
    section = st.selectbox("Section to polish", choices)
    intensity = st.selectbox("Polish intensity", ["Light", "Medium", "Deep"], index=1)
    polish_targets = st.multiselect(
        "Polish targets",
        ["Remove translation tone", "Reduce repetitive sentence patterns", "Avoid generic AI phrases", "Make dialogue less formal", "Add subtext", "Use natural Chinese rhythm", "Avoid too much exposition", "Strengthen hook", "Improve sensory detail"],
        default=["Avoid generic AI phrases", "Use natural Chinese rhythm", "Make dialogue less formal", "Avoid too much exposition"],
    )
    extra = st.text_area("Extra polish instruction", height=90, placeholder="例如：更像女频仙侠，情绪克制但张力更强。")

    if section == "Full draft excerpt":
        text_to_polish = st.session_state.full_draft[:5000]
    else:
        idx = next((i for i, ch in enumerate(st.session_state.chapters) if chapter_heading(ch) == section), 0)
        text_to_polish = st.session_state.chapter_drafts.get(idx, st.session_state.full_draft[:5000])

    if st.button("Polish ✨", type="primary"):
        with st.spinner("Polishing..."):
            placeholder = st.empty()
            result = generate_text(
                system="You are a master literary editor specializing in natural Chinese and English fiction prose.",
                user=f"""Writing context:
{writing_context()}

Section: {section}
Intensity: {intensity}
Targets: {', '.join(polish_targets)}
Extra instruction: {extra or 'None'}

Original:
{text_to_polish}

Return:
## What changed
[brief notes]

## Polished version
[rewritten text]
""",
                placeholder=placeholder,
                max_tokens=5000,
            )
        if result:
            st.session_state.polished_draft = result
            cloud_auto_save("polish")
            st.success("Polish complete.")

    if st.session_state.polished_draft:
        st.markdown(st.session_state.polished_draft)
        revision = st.text_area("Revise polished version instruction", height=80)
        if st.button("Revise polished version 🔁"):
            if revision.strip():
                with st.spinner("Revising polished version..."):
                    placeholder = st.empty()
                    result = generate_text(
                        system="You are a literary editor. Revise the polished version based on the user's instruction.",
                        user=f"Writing context:\n{writing_context()}\n\nCurrent polished result:\n{st.session_state.polished_draft}\n\nRevision instruction:\n{revision}",
                        placeholder=placeholder,
                        max_tokens=5000,
                    )
                if result:
                    st.session_state.polished_draft = result
                    cloud_auto_save("polish revision")
                    st.rerun()

    if st.button("Next: Export →", type="primary"):
        st.session_state.stage = 10
        st.rerun()

# =============================================================================
# Stage 10: Export
# =============================================================================
elif st.session_state.stage == 10:
    st.markdown('<div class="stage-header">📥 Export / 导出</div>', unsafe_allow_html=True)
    st.markdown('<div class="stage-desc">Export manuscript, Story Bible, notes, DOCX, TXT, Markdown, and JSON project backup.</div>', unsafe_allow_html=True)

    rebuild_full_draft()
    final = st.session_state.polished_draft or st.session_state.full_draft or st.session_state.outline
    word_count = len(final.split())
    char_count = len(final)
    ch_count = len(st.session_state.chapter_drafts)
    c1, c2, c3, c4 = st.columns(4)
    c1.metric("Drafted chapters", ch_count)
    c2.metric("Approx. words", f"{word_count:,}")
    c3.metric("Characters", f"{char_count:,}")
    c4.metric("Language", st.session_state.language)

    def manuscript_only_text(use_polished: bool = False) -> str:
        """Return only user-facing novel prose, without Story Bible, outline, notes, or reports."""
        if use_polished and st.session_state.polished_draft:
            return st.session_state.polished_draft.strip()
        ordered = []
        for zero_idx, chapter in enumerate(st.session_state.chapters):
            draft = (st.session_state.chapter_drafts or {}).get(zero_idx, "")
            if draft and draft.strip():
                one_idx = zero_idx + 1
                title = chapter.get("title", f"第{one_idx}章") if isinstance(chapter, dict) else f"第{one_idx}章"
                ordered.append(f"{title}\n\n{draft.strip()}")
        if ordered:
            return "\n\n".join(ordered).strip()
        return (st.session_state.full_draft or st.session_state.polished_draft or "").strip()

    def build_markdown(export_mode: str = "project", include_notes: bool = True) -> str:
        manuscript = manuscript_only_text(use_polished=(export_mode == "polished_manuscript"))
        if export_mode in {"submission", "manuscript", "polished_manuscript"}:
            return "\n\n".join([f"# {st.session_state.novel_title}", manuscript]).strip()

        parts = [
            f"# {st.session_state.novel_title}",
            f"- Language: {st.session_state.language}",
            f"- Audience: {st.session_state.audience_channel}",
            f"- Genre: {st.session_state.genre}",
            f"- Genre tags: {', '.join(st.session_state.genre_tags or [])}",
            f"- Custom tags: {st.session_state.custom_genre_tags or 'None'}",
            f"- Style: {st.session_state.chinese_style}",
            f"- Publishing platforms: {', '.join(st.session_state.publishing_platforms or [])}",
            f"- Submission types: {', '.join(st.session_state.submission_types or [])}",
            f"- Platform goals: {', '.join(st.session_state.platform_goals or [])}",
            "\n---\n",
            "# Story Bible",
            compact_story_bible(20000),
            "\n---\n",
            "# Outline",
            st.session_state.outline,
            "\n---\n",
            "# Full Draft",
            st.session_state.full_draft,
        ]
        if include_notes and st.session_state.market_reviews:
            parts += ["\n---\n", "# Web Novel Score Reports", "\n\n".join(v.get("report", "") for v in st.session_state.market_reviews.values())]
        if include_notes and st.session_state.logic_report:
            parts += ["\n---\n", "# Consistency Report", st.session_state.logic_report]
        if include_notes and st.session_state.polished_draft:
            parts += ["\n---\n", "# Polished Sections", st.session_state.polished_draft]
        return "\n\n".join(p for p in parts if p)

    def build_docx(export_mode: str = "project") -> bytes:
        doc = Document()
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_para.add_run(st.session_state.novel_title.upper())
        run.bold = True
        run.font.size = Pt(24)

        if export_mode in {"submission", "manuscript", "polished_manuscript"}:
            doc.add_page_break()
            manuscript = manuscript_only_text(use_polished=(export_mode == "polished_manuscript"))
            for line in manuscript.splitlines():
                stripped = line.strip()
                if not stripped:
                    continue
                if stripped.startswith("第") and ("章" in stripped[:8] or "章" in stripped[:12]) and len(stripped) <= 40:
                    doc.add_heading(stripped, level=1)
                elif stripped.startswith("# "):
                    doc.add_heading(stripped[2:], level=1)
                elif stripped.startswith("## "):
                    doc.add_heading(stripped[3:], level=2)
                else:
                    doc.add_paragraph(stripped)
            buf = io.BytesIO()
            doc.save(buf)
            buf.seek(0)
            return buf.getvalue()

        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub.add_run(f"Language: {st.session_state.language} | Genre: {st.session_state.genre}")
        doc.add_page_break()
        for heading, text in [
            ("Story Bible", compact_story_bible(20000)),
            ("Outline", st.session_state.outline),
            ("Full Draft", st.session_state.full_draft),
            ("Web Novel Score Reports", "\n\n".join(v.get("report", "") for v in st.session_state.market_reviews.values()) if st.session_state.market_reviews else ""),
            ("Consistency Report", st.session_state.logic_report),
            ("Polished Sections", st.session_state.polished_draft),
        ]:
            if text:
                doc.add_heading(heading, level=1)
                for line in text.splitlines():
                    if line.startswith("# "):
                        doc.add_heading(line[2:], level=1)
                    elif line.startswith("## "):
                        doc.add_heading(line[3:], level=2)
                    elif line.strip():
                        doc.add_paragraph(line.strip())
                doc.add_page_break()
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.getvalue()

    filename = safe_filename(st.session_state.project_name or st.session_state.novel_title)
    json_backup = json.dumps(project_data(), ensure_ascii=False, indent=2)

    st.subheader("Download options / 下载选项")
    export_choice = st.radio(
        "Choose export type",
        options=[
            "投稿版本：只含小说正文",
            "正文版本：小说正文 + 章节标题",
            "润色版本：只含润色正文",
            "项目完整包：设定 + 大纲 + 正文 + 报告",
            "项目备份 JSON",
        ],
        index=0,
        horizontal=False,
        key="export_type_choice_v19",
    )

    mode_map = {
        "投稿版本：只含小说正文": "submission",
        "正文版本：小说正文 + 章节标题": "manuscript",
        "润色版本：只含润色正文": "polished_manuscript",
        "项目完整包：设定 + 大纲 + 正文 + 报告": "project",
    }

    if export_choice == "项目备份 JSON":
        st.info("Backup file for restoring the whole project. It includes settings, Story Bible, outline, drafts, reviews, and cloud-save data.")
        st.download_button("Download Project JSON", json_backup, f"{filename}_project_backup.json", "application/json", use_container_width=True)
        preview_text = json_backup[:2500]
    else:
        export_mode = mode_map[export_choice]
        md = build_markdown(export_mode=export_mode)
        txt = md.replace("#", "").strip()
        suffix_map = {
            "submission": "submission_manuscript",
            "manuscript": "manuscript",
            "polished_manuscript": "polished_manuscript",
            "project": "full_project",
        }
        suffix = suffix_map[export_mode]
        if export_mode in {"submission", "manuscript", "polished_manuscript"}:
            st.success("This export contains only the novel manuscript. It does not include Story Bible, outline, beat sheets, reviews, or internal notes.")
        else:
            st.info("This export is the full working package and includes Story Bible, outline, manuscript, reviews, and reports.")

        col_a, col_b, col_c = st.columns(3)
        with col_a:
            st.download_button("Download TXT", txt, f"{filename}_{suffix}.txt", "text/plain", use_container_width=True)
        with col_b:
            st.download_button("Download Markdown", md, f"{filename}_{suffix}.md", "text/markdown", use_container_width=True)
        with col_c:
            st.download_button("Download DOCX", build_docx(export_mode=export_mode), f"{filename}_{suffix}.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)
        preview_text = txt[:2500]

    st.divider()
    st.markdown("**Preview**")
    st.text(preview_text + ("..." if len(preview_text) > 2500 else ""))
